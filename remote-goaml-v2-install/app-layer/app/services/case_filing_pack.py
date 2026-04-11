"""
Case filing evidence pack generation for the Command Center.
"""

from __future__ import annotations

from io import BytesIO
import json
from datetime import datetime, timezone
from typing import Any
from uuid import UUID

from docx import Document
from docx.shared import Inches
from fastapi.encoders import jsonable_encoder
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from core.database import get_pool
from services.case_context import get_case_context
from services.case_grounding import build_case_grounding
from services.case_workspace import get_case_filing_readiness, get_case_workflow_state
from services.cases import get_case_detail, get_case_sar, list_case_notes, list_case_tasks


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _normalize_json_dict(value: Any) -> dict[str, Any]:
    if isinstance(value, dict):
        return value
    if isinstance(value, str) and value.strip():
        try:
            parsed = json.loads(value)
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            return {}
    return {}


def _dedupe_key(item: dict[str, Any]) -> str:
    return "|".join(
        [
            str(item.get("evidence_type") or ""),
            str(item.get("source_evidence_id") or ""),
            str(item.get("title") or ""),
        ]
    )


def filing_pack_filename(case_ref: str, export_format: str) -> str:
    safe_ref = "".join(ch if ch.isalnum() or ch in {"-", "_"} else "-" for ch in str(case_ref or "case"))
    return f"{safe_ref.lower()}-filing-pack.{export_format}"


def _json_default(value: Any) -> Any:
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, UUID):
        return str(value)
    return str(value)


def _pack_to_json_bytes(pack: dict[str, Any]) -> bytes:
    encoded = jsonable_encoder(pack)
    return json.dumps(encoded, indent=2, default=_json_default).encode("utf-8")


def _evidence_table_rows(items: list[dict[str, Any]]) -> list[list[str]]:
    rows = [["Type", "Title", "Source", "SAR", "Importance"]]
    for item in items:
        rows.append(
            [
                str(item.get("evidence_type") or ""),
                str(item.get("title") or ""),
                str(item.get("source") or ""),
                "Yes" if item.get("include_in_sar") else "No",
                str(item.get("importance") or ""),
            ]
        )
    return rows


def _pack_to_pdf_bytes(pack: dict[str, Any]) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=0.6 * inch,
        rightMargin=0.6 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle("PackBody", parent=styles["BodyText"], fontSize=9.5, leading=13)
    heading = styles["Heading2"]
    story: list[Any] = []

    def add_heading(text: str) -> None:
        story.append(Paragraph(text, heading))
        story.append(Spacer(1, 0.12 * inch))

    def add_paragraph(text: str | None) -> None:
        if text:
            story.append(Paragraph(str(text).replace("\n", "<br/>"), body))
            story.append(Spacer(1, 0.08 * inch))

    story.append(Paragraph(f"Case Filing Pack - {pack.get('case_ref')}", styles["Title"]))
    story.append(Spacer(1, 0.14 * inch))
    add_paragraph(
        f"Generated {pack.get('generated_at')} by {pack.get('generated_by') or 'system'}"
    )
    add_paragraph(f"Grounding mode: {pack.get('grounding_mode') or 'context_fallback'}")

    add_heading("Executive Summary")
    for item in pack.get("summary", []) or []:
        add_paragraph(f"• {item}")

    if pack.get("ai_summary"):
        add_heading("AI Summary")
        add_paragraph(pack.get("ai_summary"))

    if pack.get("risk_factors"):
        add_heading("Risk Factors")
        for factor in pack.get("risk_factors", []):
            add_paragraph(f"• {factor}")

    readiness = pack.get("filing_readiness") or {}
    add_heading("Filing Readiness")
    readiness_lines = [
        f"Overall status: {readiness.get('overall_status')}",
        f"Score: {readiness.get('score')}",
    ]
    for item in readiness.get("blocking_items", []) or []:
        readiness_lines.append(f"Blocker: {item}")
    for item in readiness.get("warning_items", []) or []:
        readiness_lines.append(f"Warning: {item}")
    for item in readiness.get("passed_checks", []) or []:
        readiness_lines.append(f"Passed: {item}")
    for line in readiness_lines:
        add_paragraph(line)

    sar = pack.get("sar") or {}
    if sar:
        add_heading("SAR Detail")
        add_paragraph(f"SAR ref: {sar.get('sar_ref')} | Status: {sar.get('status')}")
        add_paragraph(f"Subject: {sar.get('subject_name') or 'Unknown'}")
        if sar.get("narrative"):
            add_paragraph(sar.get("narrative"))

    for section_title, items in (
        ("Filing Evidence", pack.get("filing_evidence", []) or []),
        ("Supporting Evidence", pack.get("supporting_evidence", []) or []),
    ):
        add_heading(section_title)
        if not items:
            add_paragraph("No evidence items available.")
            continue
        table = Table(_evidence_table_rows(items), repeatRows=1, colWidths=[1.1 * inch, 2.75 * inch, 1.3 * inch, 0.55 * inch, 0.8 * inch])
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eefb")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1f2a44")),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#d0d8e8")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("LEADING", (0, 0), (-1, -1), 10),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7f9fc")]),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 0.15 * inch))

    if pack.get("notes"):
        add_heading("Analyst Notes")
        for note in pack.get("notes", []):
            add_paragraph(f"{note.get('created_at')} - {note.get('author') or 'analyst'}: {note.get('text')}")

    if pack.get("tasks"):
        add_heading("Tasks")
        for task in pack.get("tasks", []):
            add_paragraph(
                f"{task.get('status')} | {task.get('priority')} | {task.get('title')} | assignee={task.get('assigned_to') or 'unassigned'}"
            )
            if task.get("description"):
                add_paragraph(task.get("description"))

    doc.build(story)
    return buffer.getvalue()


def _pack_to_docx_bytes(pack: dict[str, Any]) -> bytes:
    document = Document()
    document.add_heading(f"Case Filing Pack - {pack.get('case_ref')}", level=0)
    document.add_paragraph(
        f"Generated {pack.get('generated_at')} by {pack.get('generated_by') or 'system'}"
    )
    document.add_paragraph(f"Grounding mode: {pack.get('grounding_mode') or 'context_fallback'}")

    document.add_heading("Executive Summary", level=1)
    for item in pack.get("summary", []) or []:
        document.add_paragraph(str(item), style="List Bullet")

    if pack.get("ai_summary"):
        document.add_heading("AI Summary", level=1)
        document.add_paragraph(str(pack.get("ai_summary")))

    if pack.get("risk_factors"):
        document.add_heading("Risk Factors", level=1)
        for factor in pack.get("risk_factors", []):
            document.add_paragraph(str(factor), style="List Bullet")

    readiness = pack.get("filing_readiness") or {}
    document.add_heading("Filing Readiness", level=1)
    document.add_paragraph(f"Overall status: {readiness.get('overall_status')}")
    document.add_paragraph(f"Score: {readiness.get('score')}")
    for label, items in (
        ("Blocking items", readiness.get("blocking_items", []) or []),
        ("Warnings", readiness.get("warning_items", []) or []),
        ("Passed checks", readiness.get("passed_checks", []) or []),
    ):
        if items:
            document.add_paragraph(label)
            for item in items:
                document.add_paragraph(str(item), style="List Bullet")

    sar = pack.get("sar") or {}
    if sar:
        document.add_heading("SAR Detail", level=1)
        document.add_paragraph(f"SAR ref: {sar.get('sar_ref')} | Status: {sar.get('status')}")
        document.add_paragraph(f"Subject: {sar.get('subject_name') or 'Unknown'}")
        if sar.get("narrative"):
            document.add_paragraph(str(sar.get("narrative")))

    for section_title, items in (
        ("Filing Evidence", pack.get("filing_evidence", []) or []),
        ("Supporting Evidence", pack.get("supporting_evidence", []) or []),
    ):
        document.add_heading(section_title, level=1)
        if not items:
            document.add_paragraph("No evidence items available.")
            continue
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = table.rows[0].cells
        for idx, title in enumerate(["Type", "Title", "Source", "SAR", "Importance"]):
            headers[idx].text = title
        for item in items:
            row = table.add_row().cells
            row[0].text = str(item.get("evidence_type") or "")
            row[1].text = str(item.get("title") or "")
            row[2].text = str(item.get("source") or "")
            row[3].text = "Yes" if item.get("include_in_sar") else "No"
            row[4].text = str(item.get("importance") or "")

    if pack.get("notes"):
        document.add_heading("Analyst Notes", level=1)
        for note in pack.get("notes", []):
            document.add_paragraph(
                f"{note.get('created_at')} - {note.get('author') or 'analyst'}: {note.get('text')}",
                style="List Bullet",
            )

    if pack.get("tasks"):
        document.add_heading("Tasks", level=1)
        for task in pack.get("tasks", []):
            document.add_paragraph(
                f"{task.get('status')} | {task.get('priority')} | {task.get('title')} | assignee={task.get('assigned_to') or 'unassigned'}",
                style="List Bullet",
            )
            if task.get("description"):
                document.add_paragraph(str(task.get("description")))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_case_filing_pack(pack: dict[str, Any], export_format: str) -> tuple[str, bytes, str]:
    format_key = str(export_format or "json").lower()
    case_ref = str(pack.get("case_ref") or "case")
    if format_key == "json":
        return filing_pack_filename(case_ref, "json"), _pack_to_json_bytes(pack), "application/json"
    if format_key == "pdf":
        return filing_pack_filename(case_ref, "pdf"), _pack_to_pdf_bytes(pack), "application/pdf"
    if format_key == "docx":
        return (
            filing_pack_filename(case_ref, "docx"),
            _pack_to_docx_bytes(pack),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    raise ValueError(f"Unsupported export format: {export_format}")


async def generate_case_filing_pack(
    case_id: UUID,
    *,
    generated_by: str | None = None,
    include_notes: bool = True,
    include_tasks: bool = True,
    include_ai_summary: bool = True,
    evidence_limit: int = 12,
) -> dict[str, Any] | None:
    case_detail = await get_case_detail(case_id)
    if not case_detail:
        return None

    context = await get_case_context(case_id, document_limit=4, related_limit=8)
    grounding = await build_case_grounding(
        case_id,
        context=context,
        prioritize_pinned_evidence=True,
        filing_only=True,
        limit=evidence_limit,
    )
    broader_grounding = await build_case_grounding(
        case_id,
        context=context,
        prioritize_pinned_evidence=True,
        filing_only=False,
        limit=evidence_limit,
    )
    readiness = await get_case_filing_readiness(case_id)
    workflow = await get_case_workflow_state(case_id)
    sar = await get_case_sar(case_id)
    notes = await list_case_notes(case_id) if include_notes else []
    tasks = await list_case_tasks(case_id) if include_tasks else []
    if context is None or readiness is None or workflow is None:
        return None

    used_evidence = grounding.get("used_evidence", [])
    filing_evidence = [item for item in used_evidence if item.get("include_in_sar")] or used_evidence[:6]
    filing_keys = {_dedupe_key(item) for item in filing_evidence}
    supporting_evidence = [
        item for item in broader_grounding.get("used_evidence", [])
        if _dedupe_key(item) not in filing_keys
    ][: max(evidence_limit - len(filing_evidence), 0)]

    ai_summary = case_detail.get("ai_summary") if include_ai_summary else None
    risk_factors = case_detail.get("ai_risk_factors") or []
    summary: list[str] = [
        f"{case_detail.get('case_ref')} is a {case_detail.get('priority')} priority case in {case_detail.get('status')} status.",
        f"Filing readiness is {readiness.get('overall_status')} with score {readiness.get('score')}.",
        f"The pack contains {len(filing_evidence)} filing evidence item{'s' if len(filing_evidence) != 1 else ''} and {len(supporting_evidence)} supporting item{'s' if len(supporting_evidence) != 1 else ''}.",
    ]
    if sar and sar.get("status"):
        summary.append(f"SAR status is {sar.get('status')} for {sar.get('sar_ref')}.")
    if workflow.get("active_process"):
        summary.append(
            f"Active workflow: {workflow['active_process'].get('workflow_label') or workflow['active_process'].get('workflow_key')}."
        )
    if readiness.get("blocking_items"):
        summary.append(f"{len(readiness['blocking_items'])} filing blocker(s) still need resolution.")

    pool = get_pool()
    async with pool.acquire() as conn:
        await conn.execute(
            """
            INSERT INTO case_events (case_id, event_type, actor, detail, metadata)
            VALUES ($1, 'filing_pack_generated', $2, $3, $4::jsonb)
            """,
            case_id,
            generated_by,
            "Filing evidence pack generated",
            json.dumps(
                {
                    "grounding_mode": grounding.get("grounding_mode"),
                    "used_evidence_count": len(used_evidence),
                    "filing_evidence_count": len(filing_evidence),
                    "supporting_evidence_count": len(supporting_evidence),
                }
            ),
        )

    return {
        "case_id": case_id,
        "case_ref": case_detail["case_ref"],
        "generated_by": generated_by,
        "generated_at": _utcnow(),
        "summary": summary,
        "ai_summary": ai_summary,
        "risk_factors": risk_factors,
        "sar": None if sar == {} else sar,
        "workflow": workflow,
        "filing_readiness": readiness,
        "grounding_mode": grounding.get("grounding_mode"),
        "used_evidence": used_evidence,
        "filing_evidence": filing_evidence,
        "supporting_evidence": supporting_evidence,
        "notes": notes or [],
        "tasks": tasks or [],
    }
