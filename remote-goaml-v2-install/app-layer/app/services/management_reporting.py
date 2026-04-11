"""
Executive and manager reporting service with exportable management packs.
"""

from __future__ import annotations

from collections import Counter, defaultdict
from datetime import datetime, timedelta, timezone
from email.message import EmailMessage
from io import BytesIO, StringIO
import csv
import json
import smtplib
from typing import Any
from uuid import UUID

from docx import Document
import httpx
from fastapi.encoders import jsonable_encoder
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from core.config import settings
from core.database import get_pool
from services.decision_quality import (
    capture_decision_quality_snapshot,
    get_decision_quality_analytics,
    get_decision_quality_drilldown,
    get_decision_quality_snapshots,
)
from services.entities import list_watchlist_entities
from services.model_monitoring import get_scorer_monitoring_summary
from services.model_registry import get_scorer_model_ops_summary, get_scorer_outcome_analytics
from services.playbook_analytics import get_playbook_analytics
from services.workflow_engine import get_workflow_overview
from services.sla_analytics import get_sar_queue_trends


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _normalize_json_dict(value: Any) -> dict[str, Any]:
    if isinstance(value, dict):
        return value
    if isinstance(value, str) and value.strip():
        try:
            parsed = json.loads(value)
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            return {}
    return {}


def _normalize_json_list(value: Any) -> list[Any]:
    if isinstance(value, list):
        return value
    if isinstance(value, tuple):
        return list(value)
    if isinstance(value, str) and value.strip():
        try:
            parsed = json.loads(value)
            if isinstance(parsed, list):
                return parsed
        except json.JSONDecodeError:
            return []
    return []


def _safe_datetime(value: Any) -> datetime | None:
    if isinstance(value, datetime):
        return value if value.tzinfo else value.replace(tzinfo=timezone.utc)
    if isinstance(value, str) and value.strip():
        try:
            parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
            return parsed if parsed.tzinfo else parsed.replace(tzinfo=timezone.utc)
        except ValueError:
            return None
    return None


def _bucket_month(value: datetime | None) -> str:
    target = value or _utcnow()
    return target.astimezone(timezone.utc).strftime("%Y-%m")


def _period_bounds(granularity: str, reference: datetime | None = None) -> tuple[datetime, datetime, str]:
    ref = (reference or _utcnow()).astimezone(timezone.utc)
    key = str(granularity or "daily").lower()
    if key == "monthly":
        start = ref.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        if start.month == 12:
            next_start = start.replace(year=start.year + 1, month=1)
        else:
            next_start = start.replace(month=start.month + 1)
        end = next_start - timedelta(microseconds=1)
        label = start.strftime("%Y-%m")
        return start, end, label
    if key == "weekly":
        start = (ref - timedelta(days=ref.weekday())).replace(hour=0, minute=0, second=0, microsecond=0)
        end = start + timedelta(days=7) - timedelta(microseconds=1)
        label = f"{start.strftime('%Y-%m-%d')} to {(start + timedelta(days=6)).strftime('%Y-%m-%d')}"
        return start, end, label
    start = ref.replace(hour=0, minute=0, second=0, microsecond=0)
    end = start + timedelta(days=1) - timedelta(microseconds=1)
    label = start.strftime("%Y-%m-%d")
    return start, end, label


def _rate(numerator: int | float, denominator: int | float) -> float:
    if not denominator:
        return 0.0
    return round(float(numerator) / float(denominator), 4)


def _titleize(value: str) -> str:
    return str(value or "").replace("_", " ").title()


def _as_text_list(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item) for item in value if str(item).strip()]
    if isinstance(value, str) and value.strip():
        return [value]
    return []


async def _fetch_reporting_rows(range_days: int, *, since: datetime | None = None, until: datetime | None = None) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    pool = get_pool()
    since_value = since or (_utcnow() - timedelta(days=range_days))
    until_value = until or _utcnow()
    async with pool.acquire() as conn:
        case_rows = await conn.fetch(
            """
            SELECT
                c.id,
                c.case_ref,
                c.status,
                c.priority,
                c.assigned_to,
                c.sar_id,
                c.created_at,
                c.updated_at,
                c.metadata,
                s.status AS sar_status,
                s.drafted_at,
                s.reviewed_at,
                s.approved_at,
                s.filed_at,
                s.filing_ref,
                ARRAY_REMOVE(ARRAY_AGG(DISTINCT a.alert_type), NULL) AS alert_types,
                ARRAY_REMOVE(ARRAY_AGG(DISTINCT a.status), NULL) AS alert_statuses,
                ARRAY_REMOVE(ARRAY_AGG(DISTINCT a.rule_id), NULL) AS alert_rule_ids,
                COALESCE(ev.event_count, 0) AS event_count,
                COALESCE(ev.has_sar_drafted, FALSE) AS has_sar_drafted_event,
                COALESCE(ev.has_sar_submitted, FALSE) AS has_sar_submitted_event,
                COALESCE(ev.has_sar_approved, FALSE) AS has_sar_approved_event,
                COALESCE(ev.has_sar_filed, FALSE) AS has_sar_filed_event,
                COALESCE(ce.evidence_count, 0) AS evidence_count,
                COALESCE(ce.included_evidence_count, 0) AS included_evidence_count
            FROM cases c
            LEFT JOIN sar_reports s ON s.id = c.sar_id
            LEFT JOIN case_alerts ca ON ca.case_id = c.id
            LEFT JOIN alerts a ON a.id = ca.alert_id
            LEFT JOIN LATERAL (
                SELECT
                    COUNT(*) AS event_count,
                    BOOL_OR(event_type = 'sar_drafted') AS has_sar_drafted,
                    BOOL_OR(event_type = 'sar_submitted_for_review') AS has_sar_submitted,
                    BOOL_OR(event_type = 'sar_approved') AS has_sar_approved,
                    BOOL_OR(event_type = 'sar_filed') AS has_sar_filed
                FROM case_events cev
                WHERE cev.case_id = c.id
            ) ev ON TRUE
            LEFT JOIN LATERAL (
                SELECT
                    COUNT(*) AS evidence_count,
                    COUNT(*) FILTER (WHERE include_in_sar = TRUE) AS included_evidence_count
                FROM case_evidence cex
                WHERE cex.case_id = c.id
            ) ce ON TRUE
            WHERE c.created_at >= $1 AND c.created_at <= $2
            GROUP BY
                c.id, s.id,
                ev.event_count, ev.has_sar_drafted, ev.has_sar_submitted, ev.has_sar_approved, ev.has_sar_filed,
                ce.evidence_count, ce.included_evidence_count
            ORDER BY c.created_at DESC
            """,
            since_value,
            until_value,
        )
        alert_rows = await conn.fetch(
            """
            SELECT
                a.id,
                a.alert_ref,
                a.alert_type,
                a.status,
                a.severity,
                a.assigned_to,
                a.created_at,
                a.metadata,
                c.metadata AS case_metadata
            FROM alerts a
            LEFT JOIN cases c ON c.id = a.case_id
            WHERE a.created_at >= $1 AND a.created_at <= $2
            ORDER BY a.created_at DESC
            """,
            since_value,
            until_value,
        )
    return [dict(row) for row in case_rows], [dict(row) for row in alert_rows]


def _case_scope(case_row: dict[str, Any]) -> dict[str, str]:
    metadata = _normalize_json_dict(case_row.get("metadata"))
    routing = _normalize_json_dict(metadata.get("routing"))
    team_key = str(routing.get("team_key") or metadata.get("team_key") or "unassigned_team")
    team_label = str(routing.get("team_label") or metadata.get("team_label") or _titleize(team_key))
    region_key = str(routing.get("region_key") or metadata.get("region_key") or "global")
    region_label = str(routing.get("region_label") or metadata.get("region_label") or _titleize(region_key))
    return {
        "team_key": team_key,
        "team_label": team_label,
        "region_key": region_key,
        "region_label": region_label,
    }


def _alert_scope(alert_row: dict[str, Any]) -> dict[str, str]:
    metadata = _normalize_json_dict(alert_row.get("metadata"))
    case_metadata = _normalize_json_dict(alert_row.get("case_metadata"))
    routing = _normalize_json_dict(metadata.get("routing") or case_metadata.get("routing"))
    team_key = str(routing.get("team_key") or metadata.get("team_key") or case_metadata.get("team_key") or "unassigned_team")
    team_label = str(routing.get("team_label") or metadata.get("team_label") or case_metadata.get("team_label") or _titleize(team_key))
    region_key = str(routing.get("region_key") or metadata.get("region_key") or case_metadata.get("region_key") or "global")
    region_label = str(routing.get("region_label") or metadata.get("region_label") or case_metadata.get("region_label") or _titleize(region_key))
    return {
        "team_key": team_key,
        "team_label": team_label,
        "region_key": region_key,
        "region_label": region_label,
    }


def _json_default(value: Any) -> Any:
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, UUID):
        return str(value)
    return str(value)


def management_report_filename(report_key: str, export_format: str) -> str:
    safe = "".join(ch if ch.isalnum() or ch in {"-", "_"} else "-" for ch in report_key)
    return f"{safe.lower()}.{export_format}"


def _extract_model_version(rule_ids: list[str] | None) -> str | None:
    for rule_id in rule_ids or []:
        value = str(rule_id or "").strip()
        if ":" in value:
            return value.rsplit(":", 1)[-1] or None
    return None


def _hours_between(start: datetime | None, end: datetime | None) -> float | None:
    if not start or not end:
        return None
    return max(0.0, round((end - start).total_seconds() / 3600.0, 2))


def _scope_metric_rows(store: dict[str, dict[str, Any]]) -> list[dict[str, Any]]:
    rows = []
    for key, values in sorted(store.items(), key=lambda item: (item[1].get("count", 0), item[0]), reverse=True):
        rows.append(
            {
                "scope_key": str(values.get("scope_key") or key),
                "scope_label": str(values.get("scope_label") or key),
                "metric": float(values.get("metric") or 0.0),
                "count": int(values.get("count") or 0),
                "secondary_metric": float(values["secondary_metric"]) if values.get("secondary_metric") is not None else None,
            }
        )
    return rows


def _default_distribution_rules() -> list[dict[str, Any]]:
    return [
        {
            "rule_key": "manager-daily-csv",
            "display_name": "Manager Daily CSV",
            "target_roles": ["manager"],
            "template_key": "manager",
            "export_format": "csv",
            "cadence": "daily",
            "enabled": True,
            "channels": ["email"],
            "recipients": [],
            "metadata": {"description": "Operational manager pack with queue, SLA, and outcome analytics."},
        },
        {
            "rule_key": "executive-weekly-pdf",
            "display_name": "Executive Weekly PDF",
            "target_roles": ["manager", "auditor", "admin"],
            "template_key": "executive",
            "export_format": "pdf",
            "cadence": "weekly",
            "enabled": True,
            "channels": ["email"],
            "recipients": [],
            "metadata": {"description": "Executive summary pack focused on posture, trend, and health reporting."},
        },
        {
            "rule_key": "compliance-weekly-docx",
            "display_name": "Compliance Weekly DOCX",
            "target_roles": ["auditor", "admin"],
            "template_key": "compliance",
            "export_format": "docx",
            "cadence": "weekly",
            "enabled": False,
            "channels": ["email"],
            "recipients": [],
            "metadata": {"description": "Compliance-focused pack with filing, workflow, and typology oversight."},
        },
        {
            "rule_key": "board-monthly-pdf",
            "display_name": "Board Monthly PDF",
            "target_roles": ["manager", "admin"],
            "template_key": "board",
            "export_format": "pdf",
            "cadence": "monthly",
            "enabled": True,
            "channels": ["email"],
            "recipients": [],
            "metadata": {"description": "Board-level monthly pack with posture, top risks, and movement commentary."},
        },
        {
            "rule_key": "board-quarterly-docx",
            "display_name": "Board Quarterly DOCX",
            "target_roles": ["manager", "admin"],
            "template_key": "board",
            "export_format": "docx",
            "cadence": "quarterly",
            "enabled": True,
            "channels": ["email"],
            "recipients": [],
            "metadata": {"description": "Quarterly board pack with strategic risk posture and improvement/decline commentary."},
        },
    ]


async def _ensure_distribution_rules() -> None:
    pool = get_pool()
    async with pool.acquire() as conn:
        for rule in _default_distribution_rules():
            await conn.execute(
                """
                INSERT INTO report_distribution_rules (
                    rule_key, display_name, target_roles, template_key, export_format,
                    cadence, enabled, channels, recipients, metadata, updated_by
                )
                VALUES ($1,$2,$3::jsonb,$4,$5,$6,$7,$8::jsonb,$9::jsonb,$10::jsonb,$11)
                ON CONFLICT (rule_key) DO NOTHING
                """,
                rule["rule_key"],
                rule["display_name"],
                json.dumps(rule["target_roles"]),
                rule["template_key"],
                rule["export_format"],
                rule["cadence"],
                bool(rule["enabled"]),
                json.dumps(rule["channels"]),
                json.dumps(rule["recipients"]),
                json.dumps(rule["metadata"]),
                "system",
            )


async def _load_reporting_snapshot_history(*, snapshot_scope: str, snapshot_granularity: str = "daily", range_days: int, limit: int = 30) -> list[dict[str, Any]]:
    pool = get_pool()
    async with pool.acquire() as conn:
        rows = await conn.fetch(
            """
            SELECT captured_at, snapshot, summary_metrics, snapshot_granularity, period_start, period_end, period_label
            FROM reporting_snapshots
            WHERE snapshot_scope = $1 AND snapshot_granularity = $2 AND range_days = $3
            ORDER BY period_end DESC NULLS LAST, captured_at DESC
            LIMIT $4
            """,
            snapshot_scope,
            snapshot_granularity,
            range_days,
            limit,
        )
    return [
        {
            "captured_at": row["captured_at"],
            "snapshot": _normalize_json_dict(row["snapshot"]),
            "summary_metrics": _normalize_json_dict(row["summary_metrics"]),
            "snapshot_granularity": row["snapshot_granularity"] or "daily",
            "period_start": row["period_start"],
            "period_end": row["period_end"],
            "period_label": row["period_label"],
        }
        for row in rows
    ]


def _derive_historical_trends_from_snapshots(history: list[dict[str, Any]], key: str) -> list[dict[str, Any]]:
    points: list[dict[str, Any]] = []
    for item in reversed(history):
        bucket = item["captured_at"].astimezone(timezone.utc).strftime("%Y-%m-%d")
        for row in item.get("snapshot", {}).get(key, []) or []:
            points.append(
                {
                    "bucket": bucket,
                    "scope_key": row.get("scope_key"),
                    "scope_label": row.get("scope_label"),
                    "value": float(row.get("value") or row.get("metric") or 0.0),
                    "count": int(row.get("count") or 0),
                    "secondary_value": float(row["secondary_value"]) if row.get("secondary_value") is not None else None,
                }
            )
    return points[-40:]


def _derive_heatmap_history_from_snapshots(history: list[dict[str, Any]]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for item in reversed(history):
        bucket = item["captured_at"].astimezone(timezone.utc).strftime("%Y-%m-%d")
        for row in item.get("snapshot", {}).get("step_heatmap_trends", []) or []:
            rows.append(
                {
                    "bucket": bucket,
                    "typology": row.get("typology"),
                    "display_name": row.get("display_name"),
                    "step_key": row.get("step_key"),
                    "step_label": row.get("step_label"),
                    "affected_case_count": int(row.get("affected_case_count") or 0),
                    "affected_case_rate": float(row.get("affected_case_rate") or 0.0),
                }
            )
    return rows[-30:]


def _build_snapshot_metrics(report: dict[str, Any]) -> dict[str, Any]:
    kpis = {str(item.get("key")): item.get("value") for item in report.get("high_level_kpis", [])}
    compliance_kpis = {str(item.get("key")): item.get("value") for item in (report.get("compliance_posture") or {}).get("kpis", [])}
    top_blocked = (report.get("playbook_effectiveness") or [{}])[0]
    top_false_positive = (report.get("false_positive_by_typology") or [{}])[0]
    return {
        "active_backlog": int(kpis.get("active_backlog") or 0),
        "breached_sars": int(kpis.get("breached_sars") or 0),
        "filed_sars": int(kpis.get("filed_sars_30d") or 0),
        "watchlist_active": int(kpis.get("watchlist_active") or 0),
        "filing_timeliness_rate": float(compliance_kpis.get("filing_timeliness_rate") or 0.0),
        "audit_trail_completeness": float(compliance_kpis.get("audit_trail_completeness") or 0.0),
        "evidence_pack_completeness": float(compliance_kpis.get("evidence_pack_completeness") or 0.0),
        "false_positive_rate": float(top_false_positive.get("metric") or 0.0),
        "blocked_case_rate": float(top_blocked.get("blocked_case_rate") or 0.0),
        "production_model_version": str(kpis.get("production_model_version") or "unknown"),
        "automation_count": int(kpis.get("workflow_automation_count") or 0),
        "top_typology": (report.get("typology_mix") or [{}])[0].get("scope_label") if report.get("typology_mix") else None,
        "top_team": (report.get("filed_sar_volume_by_team") or [{}])[0].get("scope_label") if report.get("filed_sar_volume_by_team") else None,
    }


def _render_template_title(template_key: str) -> str:
    return {
        "executive": "goAML-v2 Executive Report",
        "compliance": "goAML-v2 Compliance Oversight Report",
        "manager": "goAML-v2 Manager Operations Report",
        "board": "goAML-v2 Board Oversight Report",
    }.get(template_key, "goAML-v2 Management Report")


async def _load_snapshot_by_id(snapshot_id: str | UUID) -> dict[str, Any] | None:
    pool = get_pool()
    async with pool.acquire() as conn:
        row = await conn.fetchrow(
            """
            SELECT id, snapshot_scope, snapshot_granularity, range_days, period_start, period_end, period_label,
                   captured_at, captured_by, source, summary_metrics, snapshot, metadata
            FROM reporting_snapshots
            WHERE id = $1
            """,
            UUID(str(snapshot_id)),
        )
    if not row:
        return None
    return {
        "id": row["id"],
        "snapshot_scope": row["snapshot_scope"],
        "snapshot_granularity": row["snapshot_granularity"] or "daily",
        "range_days": row["range_days"],
        "period_start": row["period_start"],
        "period_end": row["period_end"],
        "period_label": row["period_label"],
        "captured_at": row["captured_at"],
        "captured_by": row["captured_by"],
        "source": row["source"],
        "summary_metrics": _normalize_json_dict(row["summary_metrics"]),
        "snapshot": _normalize_json_dict(row["snapshot"]),
        "metadata": _normalize_json_dict(row["metadata"]),
    }


def _compute_period_over_period(points: list[dict[str, Any]]) -> list[dict[str, Any]]:
    if len(points) < 2:
        return []
    latest = points[0]
    previous = points[1]
    keys = [
        ("active_backlog", "Active backlog"),
        ("breached_sars", "Breached SARs"),
        ("filed_sars", "Filed SARs"),
        ("watchlist_active", "Watchlist active"),
        ("filing_timeliness_rate", "Timely filing rate"),
        ("false_positive_rate", "False-positive rate"),
        ("blocked_case_rate", "Blocked case rate"),
    ]
    rows: list[dict[str, Any]] = []
    for key, label in keys:
        latest_value = float(latest.get("summary_metrics", {}).get(key) or 0)
        previous_value = float(previous.get("summary_metrics", {}).get(key) or 0)
        delta = latest_value - previous_value
        delta_pct = round((delta / previous_value) * 100, 2) if previous_value else None
        rows.append(
            {
                "key": key,
                "label": label,
                "latest_period": latest.get("period_label") or latest.get("captured_at"),
                "previous_period": previous.get("period_label") or previous.get("captured_at"),
                "latest_value": latest_value,
                "previous_value": previous_value,
                "delta": delta,
                "delta_pct": delta_pct,
                "status": "up" if delta > 0 else "down" if delta < 0 else "flat",
            }
        )
    return rows


def _default_reporting_automation_settings() -> dict[str, Any]:
    return {
        "backlog_delta_warning_pct": 10.0,
        "backlog_delta_critical_pct": 20.0,
        "filing_timeliness_warning_drop_pct": 5.0,
        "filing_timeliness_critical_drop_pct": 10.0,
        "false_positive_warning_rate": 0.35,
        "false_positive_critical_rate": 0.5,
        "blocked_step_warning_rate": 0.3,
        "blocked_step_critical_rate": 0.45,
    }


async def get_reporting_automation_settings() -> dict[str, Any]:
    defaults = _default_reporting_automation_settings()
    pool = get_pool()
    async with pool.acquire() as conn:
        row = await conn.fetchrow(
            """
            SELECT config, updated_by, updated_at
            FROM app_runtime_settings
            WHERE setting_key = 'reporting_automation'
            """
        )
    config = _normalize_json_dict(row.get("config")) if row else {}
    payload = {
        **defaults,
        **{key: config.get(key, defaults[key]) for key in defaults},
        "updated_by": row.get("updated_by") if row else None,
        "updated_at": row.get("updated_at") if row else None,
    }
    payload["summary"] = [
        f"Backlog warnings trigger at {payload['backlog_delta_warning_pct']}% and critical at {payload['backlog_delta_critical_pct']}%.",
        f"Filing-timeliness drop warnings trigger at {payload['filing_timeliness_warning_drop_pct']}% and critical at {payload['filing_timeliness_critical_drop_pct']}%.",
        f"False-positive warnings trigger at {round(payload['false_positive_warning_rate'] * 100, 1)}% and blocked-step warnings at {round(payload['blocked_step_warning_rate'] * 100, 1)}%.",
    ]
    return payload


async def update_reporting_automation_settings(*, actor: str, updates: dict[str, Any]) -> dict[str, Any]:
    current = await get_reporting_automation_settings()
    merged = {
        "backlog_delta_warning_pct": float(updates.get("backlog_delta_warning_pct", current["backlog_delta_warning_pct"])),
        "backlog_delta_critical_pct": float(updates.get("backlog_delta_critical_pct", current["backlog_delta_critical_pct"])),
        "filing_timeliness_warning_drop_pct": float(updates.get("filing_timeliness_warning_drop_pct", current["filing_timeliness_warning_drop_pct"])),
        "filing_timeliness_critical_drop_pct": float(updates.get("filing_timeliness_critical_drop_pct", current["filing_timeliness_critical_drop_pct"])),
        "false_positive_warning_rate": float(updates.get("false_positive_warning_rate", current["false_positive_warning_rate"])),
        "false_positive_critical_rate": float(updates.get("false_positive_critical_rate", current["false_positive_critical_rate"])),
        "blocked_step_warning_rate": float(updates.get("blocked_step_warning_rate", current["blocked_step_warning_rate"])),
        "blocked_step_critical_rate": float(updates.get("blocked_step_critical_rate", current["blocked_step_critical_rate"])),
    }
    pool = get_pool()
    async with pool.acquire() as conn:
        await conn.execute(
            """
            INSERT INTO app_runtime_settings (setting_key, config, updated_by)
            VALUES ('reporting_automation', $1::jsonb, $2)
            ON CONFLICT (setting_key) DO UPDATE
            SET config = EXCLUDED.config,
                updated_by = EXCLUDED.updated_by,
                updated_at = NOW()
            """,
            json.dumps(merged),
            actor,
        )
    return await get_reporting_automation_settings()


def _evaluate_reporting_alerts(
    *,
    report: dict[str, Any],
    period_over_period: list[dict[str, Any]],
    settings_payload: dict[str, Any],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    period_map = {str(item.get("key")): item for item in period_over_period}
    alerts: list[dict[str, Any]] = []
    recommendations: list[dict[str, Any]] = []

    backlog = period_map.get("active_backlog")
    if backlog and backlog.get("delta_pct") is not None:
        backlog_delta = float(backlog["delta_pct"])
        if backlog_delta >= float(settings_payload["backlog_delta_warning_pct"]):
            severity = "critical" if backlog_delta >= float(settings_payload["backlog_delta_critical_pct"]) else "warning"
            alerts.append(
                {
                    "alert_key": "backlog_delta",
                    "title": "Backlog pressure increased",
                    "severity": severity,
                    "status": "triggered",
                    "metric_key": "active_backlog",
                    "metric_label": "Active backlog",
                    "current_value": backlog.get("latest_value"),
                    "threshold_value": settings_payload["backlog_delta_warning_pct"],
                    "delta_value": backlog.get("delta"),
                    "delta_pct": backlog_delta,
                    "message": f"Active backlog is up {backlog_delta}% versus the prior period.",
                    "recommendation_key": "rebalance_queue",
                    "metadata": {"previous_value": backlog.get("previous_value")},
                }
            )
            recommendations.append(
                {
                    "recommendation_key": "rebalance_queue",
                    "title": "Rebalance the busiest queue",
                    "priority": "critical" if severity == "critical" else "high",
                    "action_type": "manager_console",
                    "rationale": f"Backlog increased by {backlog_delta}% compared with the prior reporting period.",
                    "target_scope": "desk",
                    "target_key": "manager-console",
                    "deep_link": "/#manager-console",
                    "metadata": {"reason": "backlog_delta", "delta_pct": backlog_delta},
                }
            )

    filing_kpi = next((item for item in (report.get("compliance_posture") or {}).get("kpis", []) if item.get("key") == "filing_timeliness_rate"), None)
    timely_value = float(filing_kpi.get("value") or 0.0) if filing_kpi else 0.0
    timely_prev = float(period_map.get("filing_timeliness_rate", {}).get("previous_value") or timely_value)
    timely_drop = max(0.0, timely_prev - timely_value)
    if filing_kpi and timely_drop >= float(settings_payload["filing_timeliness_warning_drop_pct"]):
        severity = "critical" if timely_drop >= float(settings_payload["filing_timeliness_critical_drop_pct"]) else "warning"
        alerts.append(
            {
                "alert_key": "filing_timeliness_drop",
                "title": "Filing timeliness declined",
                "severity": severity,
                "status": "triggered",
                "metric_key": "filing_timeliness_rate",
                "metric_label": "Timely filing rate",
                "current_value": timely_value,
                "threshold_value": settings_payload["filing_timeliness_warning_drop_pct"],
                "delta_value": -round(timely_drop, 1),
                "delta_pct": round(-timely_drop, 1),
                "message": f"Timely filing rate fell by {round(timely_drop, 1)} percentage points.",
                "recommendation_key": "review_filing_flow",
                "metadata": {"previous_value": timely_prev},
            }
        )
        recommendations.append(
            {
                "recommendation_key": "review_filing_flow",
                "title": "Inspect filing-readiness blockers",
                "priority": "high",
                "action_type": "reporting_drilldown",
                "rationale": "Timely filing has slipped versus the prior period; review lag and approval lag should be checked next.",
                "target_scope": "metric",
                "target_key": "filing_timeliness",
                "deep_link": "/#reports",
                "metadata": {"metric": "filing_timeliness"},
            }
        )

    false_positive_rows = report.get("false_positive_by_typology") or []
    if false_positive_rows:
        top_fp = false_positive_rows[0]
        top_fp_rate = float(top_fp.get("metric") or 0.0)
        if top_fp_rate >= float(settings_payload["false_positive_warning_rate"]):
            severity = "critical" if top_fp_rate >= float(settings_payload["false_positive_critical_rate"]) else "warning"
            alerts.append(
                {
                    "alert_key": "false_positive_typology_spike",
                    "title": "False-positive concentration by typology",
                    "severity": severity,
                    "status": "triggered",
                    "metric_key": "false_positive_rate",
                    "metric_label": "False-positive rate",
                    "current_value": round(top_fp_rate * 100, 1),
                    "threshold_value": round(float(settings_payload["false_positive_warning_rate"]) * 100, 1),
                    "delta_value": None,
                    "delta_pct": None,
                    "message": f"{top_fp.get('scope_label')} is running at {round(top_fp_rate * 100, 1)}% false positives.",
                    "recommendation_key": "tune_typology_playbook",
                    "metadata": {"typology": top_fp.get("scope_key"), "count": top_fp.get("count")},
                }
            )
            recommendations.append(
                {
                    "recommendation_key": "tune_typology_playbook",
                    "title": "Tune noisy typology controls",
                    "priority": "high",
                    "action_type": "playbook_review",
                    "rationale": f"{top_fp.get('scope_label')} has the highest false-positive rate in the reporting window.",
                    "target_scope": "typology",
                    "target_key": top_fp.get("scope_key"),
                    "deep_link": "/#manager-console",
                    "metadata": {"typology": top_fp.get("scope_key")},
                }
            )

    playbook_effectiveness = report.get("playbook_effectiveness") or []
    if playbook_effectiveness:
        worst_blocked = max(playbook_effectiveness, key=lambda item: float(item.get("blocked_case_rate") or 0.0))
        blocked_rate = float(worst_blocked.get("blocked_case_rate") or 0.0)
        if blocked_rate >= float(settings_payload["blocked_step_warning_rate"]):
            severity = "critical" if blocked_rate >= float(settings_payload["blocked_step_critical_rate"]) else "warning"
            alerts.append(
                {
                    "alert_key": "playbook_blocked_rate",
                    "title": "Playbook blockers are accumulating",
                    "severity": severity,
                    "status": "triggered",
                    "metric_key": "blocked_case_rate",
                    "metric_label": "Blocked case rate",
                    "current_value": round(blocked_rate * 100, 1),
                    "threshold_value": round(float(settings_payload["blocked_step_warning_rate"]) * 100, 1),
                    "delta_value": None,
                    "delta_pct": None,
                    "message": f"{worst_blocked.get('display_name')} / {worst_blocked.get('priority')} cases are blocked at {round(blocked_rate * 100, 1)}%.",
                    "recommendation_key": "tune_playbook_thresholds",
                    "metadata": {"typology": worst_blocked.get("typology"), "priority": worst_blocked.get("priority")},
                }
            )
            recommendations.append(
                {
                    "recommendation_key": "tune_playbook_thresholds",
                    "title": "Tune playbook intervention thresholds",
                    "priority": "medium" if severity == "warning" else "high",
                    "action_type": "manager_console",
                    "rationale": "Blocked checklist rates suggest the current intervention thresholds or checklist expectations need review.",
                    "target_scope": "desk",
                    "target_key": "manager-console",
                    "deep_link": "/#manager-console",
                    "metadata": {"typology": worst_blocked.get("typology"), "priority": worst_blocked.get("priority")},
                }
            )

    deduped: dict[str, dict[str, Any]] = {}
    for item in recommendations:
        deduped.setdefault(str(item["recommendation_key"]), item)
    return alerts, list(deduped.values())


def _build_board_reporting(*, report: dict[str, Any], period_over_period: list[dict[str, Any]]) -> dict[str, Any]:
    improvements = sorted(
        [item for item in period_over_period if item.get("delta_pct") is not None and float(item.get("delta_pct") or 0) < 0],
        key=lambda item: float(item.get("delta_pct") or 0),
    )[:3]
    declines = sorted(
        [item for item in period_over_period if item.get("delta_pct") is not None and float(item.get("delta_pct") or 0) > 0],
        key=lambda item: float(item.get("delta_pct") or 0),
        reverse=True,
    )[:3]
    top_typologies = (report.get("typology_mix") or [])[:5]
    top_risks = (report.get("reporting_alerts") or [])[:5]
    commentary: list[str] = []
    for item in declines[:2]:
        commentary.append(f"{item.get('label')} worsened by {item.get('delta_pct')}% compared with the prior period.")
    for item in improvements[:2]:
        commentary.append(f"{item.get('label')} improved by {abs(float(item.get('delta_pct') or 0))}% compared with the prior period.")
    if not commentary:
        commentary.append("Board posture remains stable versus the prior comparable period.")
    return {
        "top_risks": top_risks,
        "top_typologies": top_typologies,
        "biggest_improvements": improvements,
        "biggest_declines": declines,
        "commentary": commentary,
        "quarterly_summary": [
            f"Top typology by case volume is {(top_typologies[0] or {}).get('scope_label') if top_typologies else 'n/a'}.",
            f"Filed SAR volume leader is {(report.get('filed_sar_volume_by_team') or [{}])[0].get('scope_label', 'n/a')}.",
            f"Primary model/workflow health summary: {' '.join(_as_text_list((report.get('model_workflow_health') or {}).get('workflow', {}).get('summary'))[:1]) or 'No workflow summary available.'}",
        ],
    }


async def _build_workflow_effectiveness(*, range_days: int, case_rows: list[dict[str, Any]]) -> dict[str, Any]:
    pool = get_pool()
    async with pool.acquire() as conn:
        event_rows = await conn.fetch(
            """
            SELECT
                DATE_TRUNC('day', created_at) AS bucket,
                event_type,
                COUNT(*)::int AS event_count,
                COUNT(DISTINCT case_id)::int AS touched_case_count
            FROM case_events
            WHERE created_at >= NOW() - make_interval(days => $1::int)
              AND event_type IN (
                'sar_workload_rebalanced',
                'playbook_stuck_flagged',
                'playbook_evidence_gap_escalated',
                'watchlist_case_escalated',
                'watchlist_rescreened'
              )
            GROUP BY 1, 2
            ORDER BY 1 ASC, 2 ASC
            """,
            range_days,
        )
        notification_rows = await conn.fetch(
            """
            SELECT notification_type, COUNT(*)::int AS count
            FROM notification_events
            WHERE created_at >= NOW() - make_interval(days => $1::int)
              AND notification_type IN (
                'playbook_stuck_case',
                'playbook_evidence_gap_escalation',
                'watchlist_rescreen_completed',
                'reporting_threshold_alert'
              )
            GROUP BY notification_type
            """,
            range_days,
        )

    workflow_definitions = {
        "sar_rebalance": {
            "title": "SAR Queue Rebalance",
            "event_types": {"sar_workload_rebalanced"},
            "positive": lambda case: str(case.get("status") or "").lower() in {"pending_review", "approved", "sar_filed"},
            "breached": lambda case: str(case.get("status") or "").lower() in {"overdue_review", "overdue_approval"},
        },
        "playbook_automation": {
            "title": "Playbook Automation",
            "event_types": {"playbook_stuck_flagged", "playbook_evidence_gap_escalated"},
            "positive": lambda case: str(case.get("status") or "").lower() in {"reviewing", "pending_review", "approved", "sar_filed"},
            "breached": lambda case: str(case.get("status") or "").lower() in {"overdue_review", "overdue_approval"},
        },
        "watchlist_rescreen": {
            "title": "Watchlist Re-screen",
            "event_types": {"watchlist_case_escalated", "watchlist_rescreened"},
            "positive": lambda case: str(_normalize_json_dict(case.get("metadata")).get("entity_workflow") or "").lower() == "watchlist_review" or str(case.get("status") or "").lower() in {"reviewing", "pending_review", "approved", "sar_filed"},
            "breached": lambda case: str(case.get("status") or "").lower() in {"overdue_review", "overdue_approval"},
        },
    }
    event_summary: dict[str, dict[str, Any]] = {
        key: {"triggered_count": 0, "touched_case_count": 0, "trend": []}
        for key in workflow_definitions
    }
    for row in event_rows:
        event_type = str(row["event_type"] or "")
        bucket = row["bucket"].isoformat() if row["bucket"] else None
        for workflow_key, definition in workflow_definitions.items():
            if event_type not in definition["event_types"]:
                continue
            event_summary[workflow_key]["triggered_count"] += int(row["event_count"] or 0)
            event_summary[workflow_key]["touched_case_count"] += int(row["touched_case_count"] or 0)
            event_summary[workflow_key]["trend"].append(
                {
                    "bucket": bucket,
                    "workflow_key": workflow_key,
                    "workflow_label": definition["title"],
                    "triggered_count": int(row["event_count"] or 0),
                    "touched_case_count": int(row["touched_case_count"] or 0),
                    "positive_outcome_rate": None,
                    "breached_rate": None,
                }
            )

    notification_counts = {str(row["notification_type"] or ""): int(row["count"] or 0) for row in notification_rows}
    items: list[dict[str, Any]] = []
    trends: list[dict[str, Any]] = []
    for workflow_key, definition in workflow_definitions.items():
        touched_cases = [
            case for case in case_rows
            if (
                workflow_key == "sar_rebalance"
                and str(_normalize_json_dict(case.get("metadata")).get("last_assignment_source") or "").lower() == "sar_rebalance"
            ) or (
                workflow_key == "playbook_automation"
                and bool(_normalize_json_dict(case.get("metadata")).get("playbook_automation"))
            ) or (
                workflow_key == "watchlist_rescreen"
                and str(_normalize_json_dict(case.get("metadata")).get("entity_workflow") or "").lower() == "watchlist_review"
            )
        ]
        positive_count = sum(1 for case in touched_cases if definition["positive"](case))
        breached_count = sum(1 for case in touched_cases if definition["breached"](case))
        avg_cycle_hours = None
        cycle_values = []
        for case in touched_cases:
            created_at = _safe_datetime(case.get("created_at"))
            updated_at = _safe_datetime(case.get("updated_at"))
            if created_at and updated_at:
                cycle_values.append(max(0.0, round((updated_at - created_at).total_seconds() / 3600.0, 2)))
        if cycle_values:
            avg_cycle_hours = round(sum(cycle_values) / len(cycle_values), 2)
        item = {
            "workflow_key": workflow_key,
            "title": definition["title"],
            "status": "warning" if breached_count else "stable",
            "triggered_count": int(event_summary[workflow_key]["triggered_count"]),
            "touched_case_count": len(touched_cases) or int(event_summary[workflow_key]["touched_case_count"]),
            "positive_outcome_rate": _rate(positive_count, len(touched_cases)),
            "breached_rate": _rate(breached_count, len(touched_cases)),
            "avg_cycle_hours": avg_cycle_hours,
            "note": (
                f"{notification_counts.get('playbook_evidence_gap_escalation', 0)} evidence-gap alerts were raised."
                if workflow_key == "playbook_automation"
                else (
                    f"{notification_counts.get('watchlist_rescreen_completed', 0)} watchlist re-screen notifications were generated."
                    if workflow_key == "watchlist_rescreen"
                    else "Tracks current rebalance impact across reassigned SAR cases."
                )
            ),
            "metadata": {
                "notification_count": (
                    notification_counts.get("playbook_stuck_case", 0) + notification_counts.get("playbook_evidence_gap_escalation", 0)
                    if workflow_key == "playbook_automation"
                    else notification_counts.get("watchlist_rescreen_completed", 0)
                    if workflow_key == "watchlist_rescreen"
                    else 0
                ),
            },
        }
        items.append(item)
        trends.extend(event_summary[workflow_key]["trend"])

    summary = [
        f"{sum(item['triggered_count'] for item in items)} workflow-effectiveness events were recorded in the last {range_days} days.",
        f"{sum(item['touched_case_count'] for item in items)} case touchpoints are represented across rebalance, playbook, and watchlist automation.",
    ]
    if items:
        top_item = max(items, key=lambda item: float(item.get("positive_outcome_rate") or 0.0))
        summary.append(
            f"{top_item['title']} currently has the strongest positive-outcome rate at {round(float(top_item.get('positive_outcome_rate') or 0.0) * 100, 1)}%."
        )
    return {
        "items": items,
        "trends": trends,
        "summary": summary,
    }


async def get_management_reporting_overview(*, range_days: int = 180, snapshot_scope: str = "manager") -> dict[str, Any]:
    case_rows, alert_rows = await _fetch_reporting_rows(range_days)
    playbook = await get_playbook_analytics(range_days=range_days, top_steps=16)
    workflow = await get_workflow_overview()
    model_monitoring = await get_scorer_monitoring_summary()
    model_ops = await get_scorer_model_ops_summary()
    scorer_outcomes = await get_scorer_outcome_analytics(days=min(range_days, 180))
    decision_quality = await get_decision_quality_analytics(range_days=min(range_days, 180))
    decision_quality_snapshots = await get_decision_quality_snapshots(snapshot_granularity="daily", range_days=min(range_days, 180), limit=30, auto_capture=True)
    decision_quality["snapshot_history"] = decision_quality_snapshots.get("points", [])
    decision_quality["period_over_period"] = decision_quality_snapshots.get("period_over_period", [])
    decision_quality["snapshot_summary"] = decision_quality_snapshots.get("summary", [])
    watchlist = await list_watchlist_entities(status="active", limit=100, offset=0)
    sla_trends = await get_sar_queue_trends(hours=24 * 30, limit=30, auto_capture=True, bootstrap_if_empty=True)
    snapshot_history = await _load_reporting_snapshot_history(snapshot_scope=snapshot_scope, snapshot_granularity="daily", range_days=range_days, limit=30)
    period_over_period = _compute_period_over_period(snapshot_history)
    workflow_effectiveness = await _build_workflow_effectiveness(range_days=range_days, case_rows=case_rows)

    typology_mix = Counter()
    monthly_case_counts = Counter()
    playbook_team_trends: dict[tuple[str, str], dict[str, Any]] = defaultdict(lambda: {"progress_sum": 0.0, "count": 0, "blocked": 0})
    playbook_region_trends: dict[tuple[str, str], dict[str, Any]] = defaultdict(lambda: {"progress_sum": 0.0, "count": 0, "blocked": 0})
    step_heatmap_trends: dict[tuple[str, str, str], dict[str, Any]] = defaultdict(lambda: {"count": 0, "affected_case_rate_sum": 0.0})
    effectiveness: dict[tuple[str, str], dict[str, Any]] = defaultdict(lambda: {"count": 0, "progress_sum": 0.0, "blocked": 0, "missing": 0, "false_positive": 0, "sar": 0, "filed": 0, "display_name": None})

    false_positive_team = defaultdict(lambda: {"count": 0, "fp": 0})
    false_positive_typology = defaultdict(lambda: {"count": 0, "fp": 0})
    false_positive_owner = defaultdict(lambda: {"count": 0, "fp": 0})
    filed_sar_by_team = Counter()
    filed_sar_by_region = Counter()
    outcome_correlations = defaultdict(
        lambda: {
            "case_count": 0,
            "progress_sum": 0.0,
            "delay_sum": 0.0,
            "delay_count": 0,
            "blocked": 0,
            "missing": 0,
            "sar": 0,
            "filed": 0,
            "false_positive": 0,
            "model_versions": Counter(),
        }
    )
    compliance_rollup = {
        "review_lag_sum": 0.0,
        "review_lag_count": 0,
        "approval_lag_sum": 0.0,
        "approval_lag_count": 0,
        "filing_lag_sum": 0.0,
        "filing_lag_count": 0,
        "timely_filing_count": 0,
        "filed_case_count": 0,
        "audit_complete_sum": 0.0,
        "audit_complete_count": 0,
        "evidence_complete_sum": 0.0,
        "evidence_complete_count": 0,
    }
    compliance_by_team = defaultdict(lambda: {"review_lag_sum": 0.0, "review_lag_count": 0, "approval_lag_sum": 0.0, "approval_lag_count": 0, "filed": 0, "timely": 0})
    compliance_by_region = defaultdict(lambda: {"filing_lag_sum": 0.0, "filing_lag_count": 0, "audit_sum": 0.0, "audit_count": 0, "evidence_sum": 0.0, "evidence_count": 0})

    for case in case_rows:
        scope = _case_scope(case)
        bucket = _bucket_month(_safe_datetime(case.get("created_at")))
        monthly_case_counts[bucket] += 1
        metadata = _normalize_json_dict(case.get("metadata"))
        playbook_meta = _normalize_json_dict(metadata.get("playbook"))
        typology = str(playbook_meta.get("typology") or metadata.get("typology") or ((case.get("alert_types") or [None])[0] or "unknown"))
        display_name = _titleize(typology)
        priority = str(case.get("priority") or "medium").lower()
        typology_mix[typology] += 1

        matching_typology = next((item for item in playbook.get("typologies", []) if item.get("typology") == typology), None)
        avg_progress = float((matching_typology or {}).get("avg_progress") or 0)
        blocked_case_rate = float((matching_typology or {}).get("blocked_case_rate") or 0)
        missing_case_rate = float((matching_typology or {}).get("missing_evidence_case_rate") or 0)
        fp_case = "false_positive" in [str(item).lower() for item in (case.get("alert_statuses") or [])]
        has_sar = bool(case.get("sar_id"))
        filed = str(case.get("status") or "").lower() == "sar_filed"
        model_version = _extract_model_version(case.get("alert_rule_ids") or [])
        created_at = _safe_datetime(case.get("created_at"))
        updated_at = _safe_datetime(case.get("updated_at"))
        drafted_at = _safe_datetime(case.get("drafted_at"))
        reviewed_at = _safe_datetime(case.get("reviewed_at"))
        approved_at = _safe_datetime(case.get("approved_at"))
        filed_at = _safe_datetime(case.get("filed_at"))
        workflow_delay_hours = None
        if created_at and updated_at:
            workflow_delay_hours = max(0.0, round((updated_at - created_at).total_seconds() / 3600.0, 2))
        review_lag_hours = _hours_between(drafted_at or created_at, reviewed_at)
        approval_lag_hours = _hours_between(reviewed_at or drafted_at or created_at, approved_at)
        filing_lag_hours = _hours_between(approved_at or reviewed_at or drafted_at or created_at, filed_at)
        evidence_count = int(case.get("evidence_count") or 0)
        included_evidence_count = int(case.get("included_evidence_count") or 0)
        required_evidence_min = 2 if has_sar or str(case.get("status") or "").lower() in {"pending_sar", "sar_filed"} else 1
        audit_expected = 1 + (1 if has_sar else 0) + (1 if str(case.get("sar_status") or "") in {"pending_review", "approved", "filed"} else 0) + (1 if str(case.get("sar_status") or "") in {"approved", "filed"} else 0) + (1 if filed else 0)
        audit_present = 1 + (1 if case.get("has_sar_drafted_event") else 0) + (1 if case.get("has_sar_submitted_event") else 0) + (1 if case.get("has_sar_approved_event") else 0) + (1 if case.get("has_sar_filed_event") else 0)
        audit_trail_completeness = round(min(1.0, audit_present / max(1, audit_expected)), 4)
        evidence_pack_completeness = round(min(1.0, included_evidence_count / max(1, required_evidence_min)), 4)
        timely_filing = 1 if filing_lag_hours is not None and filing_lag_hours <= 48 else 0

        team_slot = playbook_team_trends[(bucket, scope["team_key"])]
        team_slot["scope_label"] = scope["team_label"]
        team_slot["progress_sum"] += avg_progress
        team_slot["count"] += 1
        team_slot["blocked"] += 1 if blocked_case_rate > 0 else 0

        region_slot = playbook_region_trends[(bucket, scope["region_key"])]
        region_slot["scope_label"] = scope["region_label"]
        region_slot["progress_sum"] += avg_progress
        region_slot["count"] += 1
        region_slot["blocked"] += 1 if blocked_case_rate > 0 else 0

        eff = effectiveness[(typology, priority)]
        eff["display_name"] = display_name
        eff["count"] += 1
        eff["progress_sum"] += avg_progress
        eff["blocked"] += 1 if blocked_case_rate > 0 else 0
        eff["missing"] += 1 if missing_case_rate > 0 else 0
        eff["false_positive"] += 1 if fp_case else 0
        eff["sar"] += 1 if has_sar else 0
        eff["filed"] += 1 if filed else 0

        if filed:
            filed_sar_by_team[scope["team_label"]] += 1
            filed_sar_by_region[scope["region_label"]] += 1

        correlation = outcome_correlations[(typology, priority)]
        correlation["display_name"] = display_name
        correlation["case_count"] += 1
        correlation["progress_sum"] += avg_progress
        correlation["blocked"] += 1 if blocked_case_rate > 0 else 0
        correlation["missing"] += 1 if missing_case_rate > 0 else 0
        correlation["sar"] += 1 if has_sar else 0
        correlation["filed"] += 1 if filed else 0
        correlation["false_positive"] += 1 if fp_case else 0
        if workflow_delay_hours is not None:
            correlation["delay_sum"] += workflow_delay_hours
            correlation["delay_count"] += 1
        if model_version:
            correlation["model_versions"][model_version] += 1

        if review_lag_hours is not None:
            compliance_rollup["review_lag_sum"] += review_lag_hours
            compliance_rollup["review_lag_count"] += 1
            compliance_by_team[scope["team_label"]]["review_lag_sum"] += review_lag_hours
            compliance_by_team[scope["team_label"]]["review_lag_count"] += 1
        if approval_lag_hours is not None:
            compliance_rollup["approval_lag_sum"] += approval_lag_hours
            compliance_rollup["approval_lag_count"] += 1
            compliance_by_team[scope["team_label"]]["approval_lag_sum"] += approval_lag_hours
            compliance_by_team[scope["team_label"]]["approval_lag_count"] += 1
        if filing_lag_hours is not None:
            compliance_rollup["filing_lag_sum"] += filing_lag_hours
            compliance_rollup["filing_lag_count"] += 1
            compliance_by_region[scope["region_label"]]["filing_lag_sum"] += filing_lag_hours
            compliance_by_region[scope["region_label"]]["filing_lag_count"] += 1
        if filed:
            compliance_rollup["filed_case_count"] += 1
            compliance_rollup["timely_filing_count"] += timely_filing
            compliance_by_team[scope["team_label"]]["filed"] += 1
            compliance_by_team[scope["team_label"]]["timely"] += timely_filing
        compliance_rollup["audit_complete_sum"] += audit_trail_completeness
        compliance_rollup["audit_complete_count"] += 1
        compliance_rollup["evidence_complete_sum"] += evidence_pack_completeness
        compliance_rollup["evidence_complete_count"] += 1
        compliance_by_region[scope["region_label"]]["audit_sum"] += audit_trail_completeness
        compliance_by_region[scope["region_label"]]["audit_count"] += 1
        compliance_by_region[scope["region_label"]]["evidence_sum"] += evidence_pack_completeness
        compliance_by_region[scope["region_label"]]["evidence_count"] += 1

    for step in playbook.get("worst_offending_steps", []) or []:
        key = (step.get("bucket") or "current", step.get("typology") or "unknown", step.get("step_key") or "step")
        slot = step_heatmap_trends[key]
        slot["display_name"] = step.get("display_name") or _titleize(step.get("typology") or "unknown")
        slot["step_label"] = step.get("step_label") or step.get("step_key") or "Step"
        slot["count"] += int(step.get("affected_case_count") or 0)
        slot["affected_case_rate_sum"] += float(step.get("affected_case_rate") or 0)

    for alert in alert_rows:
        scope = _alert_scope(alert)
        typology = str(alert.get("alert_type") or "unknown")
        owner = str(alert.get("assigned_to") or "unassigned")
        is_fp = str(alert.get("status") or "").lower() == "false_positive"
        for key, store in (
            (scope["team_label"], false_positive_team),
            (_titleize(typology), false_positive_typology),
            (owner, false_positive_owner),
        ):
            store[key]["count"] += 1
            store[key]["fp"] += 1 if is_fp else 0

    high_level_kpis = [
        {
            "key": "active_backlog",
            "label": "Active backlog",
            "value": int(workflow.get("counts", {}).get("sar_breached", 0) + workflow.get("counts", {}).get("sar_due_soon", 0) + len(case_rows)),
            "note": "Cases plus near-term SAR queue pressure",
            "status": "info",
        },
        {
            "key": "breached_sars",
            "label": "Breached SARs",
            "value": int(workflow.get("counts", {}).get("sar_breached", 0)),
            "note": "Current breached review/approval work",
            "status": "warning" if int(workflow.get("counts", {}).get("sar_breached", 0)) else "stable",
        },
        {
            "key": "filed_sars_30d",
            "label": "Filed SARs",
            "value": int(sum(filed_sar_by_team.values())),
            "note": "Filed during reporting window",
            "status": "info",
        },
        {
            "key": "watchlist_active",
            "label": "Active watchlist entities",
            "value": int((watchlist.get("counts") or {}).get("active", 0)),
            "note": f"{(watchlist.get('counts') or {}).get('due_for_rescreen', 0)} due for re-screen",
            "status": "info",
        },
        {
            "key": "production_model_version",
            "label": "Production scorer version",
            "value": str((model_ops.get("runtime") or {}).get("model_version") or (model_ops.get("production_version") or "unknown")),
            "note": str((model_ops.get("runtime") or {}).get("model_stage") or "unknown"),
            "status": str((model_monitoring.get("latest_drift_observation") or {}).get("severity") or "stable"),
        },
        {
            "key": "workflow_automation_count",
            "label": "Active automations",
            "value": int((workflow.get("n8n") or {}).get("counts", {}).get("active_workflow_count", 0)),
            "note": f"{(workflow.get('camunda') or {}).get('counts', {}).get('tracked_process_count', 0)} tracked Camunda processes",
            "status": "info",
        },
    ]

    team_trends = [
        {
            "bucket": bucket,
            "scope_key": scope_key,
            "scope_label": values["scope_label"],
            "value": round(values["progress_sum"] / values["count"], 1) if values["count"] else 0.0,
            "count": values["count"],
            "secondary_value": _rate(values["blocked"], values["count"]),
        }
        for (bucket, scope_key), values in sorted(playbook_team_trends.items(), key=lambda item: item[0])
    ]
    region_trends = [
        {
            "bucket": bucket,
            "scope_key": scope_key,
            "scope_label": values["scope_label"],
            "value": round(values["progress_sum"] / values["count"], 1) if values["count"] else 0.0,
            "count": values["count"],
            "secondary_value": _rate(values["blocked"], values["count"]),
        }
        for (bucket, scope_key), values in sorted(playbook_region_trends.items(), key=lambda item: item[0])
    ]
    heatmap_trends = [
        {
            "bucket": bucket,
            "typology": typology,
            "display_name": values["display_name"],
            "step_key": step_key,
            "step_label": values["step_label"],
            "affected_case_count": values["count"],
            "affected_case_rate": round(values["affected_case_rate_sum"] / values["count"], 4) if values["count"] else 0.0,
        }
        for (bucket, typology, step_key), values in sorted(step_heatmap_trends.items(), key=lambda item: (item[0][0], values["count"] if False else 0))
    ]
    heatmap_trends = sorted(heatmap_trends, key=lambda item: (item["bucket"], item["affected_case_count"]), reverse=True)[:18]
    if snapshot_history:
        team_trends = _derive_historical_trends_from_snapshots(snapshot_history, "team_playbook_trends") or team_trends
        region_trends = _derive_historical_trends_from_snapshots(snapshot_history, "region_playbook_trends") or region_trends
        heatmap_trends = _derive_heatmap_history_from_snapshots(snapshot_history) or heatmap_trends

    playbook_effectiveness = [
        {
            "typology": typology,
            "display_name": values["display_name"],
            "priority": priority,
            "case_count": values["count"],
            "avg_progress": round(values["progress_sum"] / values["count"], 1) if values["count"] else 0.0,
            "blocked_case_rate": _rate(values["blocked"], values["count"]),
            "missing_evidence_case_rate": _rate(values["missing"], values["count"]),
            "false_positive_rate": _rate(values["false_positive"], values["count"]),
            "sar_conversion_rate": _rate(values["sar"], values["count"]),
            "filed_sar_rate": _rate(values["filed"], values["count"]),
        }
        for (typology, priority), values in sorted(effectiveness.items(), key=lambda item: (item[1]["count"], item[1]["sar"]), reverse=True)
    ]

    outcome_correlation_rows = [
        {
            "typology": typology,
            "display_name": values["display_name"],
            "priority": priority,
            "case_count": values["case_count"],
            "avg_progress": round(values["progress_sum"] / values["case_count"], 1) if values["case_count"] else 0.0,
            "avg_workflow_delay_hours": round(values["delay_sum"] / values["delay_count"], 2) if values["delay_count"] else None,
            "blocked_case_rate": _rate(values["blocked"], values["case_count"]),
            "missing_evidence_case_rate": _rate(values["missing"], values["case_count"]),
            "false_positive_rate": _rate(values["false_positive"], values["case_count"]),
            "sar_conversion_rate": _rate(values["sar"], values["case_count"]),
            "filed_sar_rate": _rate(values["filed"], values["case_count"]),
            "dominant_model_version": values["model_versions"].most_common(1)[0][0] if values["model_versions"] else None,
            "model_version_mix": [{ "version": version, "count": count } for version, count in values["model_versions"].most_common(3)],
        }
        for (typology, priority), values in sorted(outcome_correlations.items(), key=lambda item: (item[1]["case_count"], item[1]["sar"]), reverse=True)
    ]

    def _breakdown_rows(store: dict[str, dict[str, Any]], *, metric_label: str = "rate") -> list[dict[str, Any]]:
        rows = []
        for key, values in sorted(store.items(), key=lambda item: item[1]["count"], reverse=True):
            rows.append(
                {
                    "scope_key": key.lower().replace(" ", "_"),
                    "scope_label": key,
                    "metric": _rate(values["fp"], values["count"]),
                    "count": values["count"],
                    "secondary_metric": float(values["fp"]),
                }
            )
        return rows[:10]

    case_to_sar_by_typology = [
        {
            "scope_key": item["typology"],
            "scope_label": item["display_name"],
            "metric": item["sar_conversion_rate"],
            "count": item["case_count"],
            "secondary_metric": item["filed_sar_rate"],
        }
        for item in playbook.get("typologies", []) or []
    ]
    filed_by_team_rows = [
        {"scope_key": key.lower().replace(" ", "_"), "scope_label": key, "metric": float(count), "count": int(count)}
        for key, count in filed_sar_by_team.most_common(10)
    ]
    filed_by_region_rows = [
        {"scope_key": key.lower().replace(" ", "_"), "scope_label": key, "metric": float(count), "count": int(count)}
        for key, count in filed_sar_by_region.most_common(10)
    ]
    backlog_trends = [
        {
            "bucket": str(point.get("captured_at")),
            "scope_key": "sar_queue",
            "scope_label": "SAR Queue",
            "value": float(point.get("overall_breached_count") or 0),
            "count": int(point.get("counts", {}).get("total", 0) if isinstance(point.get("counts"), dict) else 0),
            "secondary_value": float(point.get("avg_active_age_hours") or 0),
        }
        for point in (sla_trends.get("points") or [])
    ]

    monthly_summary = [
        f"Reporting window covers {range_days} days with {len(case_rows)} cases and {len(alert_rows)} alerts.",
        f"Top typology by volume: {_titleize(typology_mix.most_common(1)[0][0]) if typology_mix else 'n/a'}.",
        f"Current watchlist posture: {(watchlist.get('counts') or {}).get('active', 0)} active, {(watchlist.get('counts') or {}).get('due_for_rescreen', 0)} due for re-screen.",
        f"Production scorer version {(model_ops.get('runtime') or {}).get('model_version') or 'unknown'} is reporting drift state {str((model_monitoring.get('latest_drift_observation') or {}).get('severity') or 'stable')}.",
    ]

    avg_review_lag = round(compliance_rollup["review_lag_sum"] / compliance_rollup["review_lag_count"], 2) if compliance_rollup["review_lag_count"] else 0.0
    avg_approval_lag = round(compliance_rollup["approval_lag_sum"] / compliance_rollup["approval_lag_count"], 2) if compliance_rollup["approval_lag_count"] else 0.0
    avg_filing_lag = round(compliance_rollup["filing_lag_sum"] / compliance_rollup["filing_lag_count"], 2) if compliance_rollup["filing_lag_count"] else 0.0
    timely_filing_rate = _rate(compliance_rollup["timely_filing_count"], compliance_rollup["filed_case_count"])
    audit_completeness_rate = _rate(compliance_rollup["audit_complete_sum"], compliance_rollup["audit_complete_count"])
    evidence_completeness_rate = _rate(compliance_rollup["evidence_complete_sum"], compliance_rollup["evidence_complete_count"])
    compliance_posture = {
        "kpis": [
            {"key": "review_lag_hours", "label": "Average review lag", "value": avg_review_lag, "status": "warning" if avg_review_lag > 24 else "stable"},
            {"key": "approval_lag_hours", "label": "Average approval lag", "value": avg_approval_lag, "status": "warning" if avg_approval_lag > 24 else "stable"},
            {"key": "filing_timeliness_rate", "label": "Timely filing rate", "value": round(timely_filing_rate * 100, 1), "status": "warning" if timely_filing_rate < 0.8 else "stable"},
            {"key": "audit_trail_completeness", "label": "Audit trail completeness", "value": round(audit_completeness_rate * 100, 1), "status": "warning" if audit_completeness_rate < 0.8 else "stable"},
            {"key": "evidence_pack_completeness", "label": "Evidence-pack completeness", "value": round(evidence_completeness_rate * 100, 1), "status": "warning" if evidence_completeness_rate < 0.75 else "stable"},
            {"key": "average_filing_lag_hours", "label": "Average filing lag", "value": avg_filing_lag, "status": "warning" if avg_filing_lag > 48 else "stable"},
        ],
        "by_team": [
            {
                "scope_key": key.lower().replace(" ", "_"),
                "scope_label": key,
                "metric": round(_rate(values["timely"], values["filed"]) * 100, 1),
                "count": int(values["filed"]),
                "secondary_metric": round(values["review_lag_sum"] / values["review_lag_count"], 2) if values["review_lag_count"] else None,
            }
            for key, values in sorted(compliance_by_team.items(), key=lambda item: item[1]["filed"], reverse=True)
        ][:8],
        "by_region": [
            {
                "scope_key": key.lower().replace(" ", "_"),
                "scope_label": key,
                "metric": round(values["audit_sum"] / values["audit_count"] * 100, 1) if values["audit_count"] else 0.0,
                "count": int(values["audit_count"]),
                "secondary_metric": round(values["evidence_sum"] / values["evidence_count"] * 100, 1) if values["evidence_count"] else None,
            }
            for key, values in sorted(compliance_by_region.items(), key=lambda item: item[1]["audit_count"], reverse=True)
        ][:8],
        "summary": [
            f"Average review lag is {avg_review_lag} hours and approval lag is {avg_approval_lag} hours.",
            f"Timely filing rate is {round(timely_filing_rate * 100, 1)}% across {compliance_rollup['filed_case_count']} filed case(s).",
            f"Audit completeness is {round(audit_completeness_rate * 100, 1)}% and evidence-pack completeness is {round(evidence_completeness_rate * 100, 1)}%.",
        ],
    }

    summary = monthly_summary + [
        f"Playbook analytics currently span {len(playbook.get('typologies', []) or [])} typologies.",
        f"False-positive rate is highest in {(_breakdown_rows(false_positive_typology) or [{'scope_label': 'n/a'}])[0]['scope_label']}.",
        f"Filed SAR volume is highest in {(filed_by_team_rows or [{'scope_label': 'n/a'}])[0]['scope_label']}.",
        f"Compliance posture shows {round(timely_filing_rate * 100, 1)}% timely filing with {round(audit_completeness_rate * 100, 1)}% audit completeness.",
    ]

    report_stub = {
        "generated_at": _utcnow(),
        "range_days": range_days,
        "snapshot_scope": snapshot_scope,
        "snapshot_granularity": "daily",
        "high_level_kpis": high_level_kpis,
        "monthly_summary": monthly_summary,
        "typology_mix": [
            {
                "scope_key": key,
                "scope_label": _titleize(key),
                "metric": float(count),
                "count": int(count),
                "secondary_metric": None,
            }
            for key, count in typology_mix.most_common(8)
        ],
        "watchlist_screening_posture": {
            "counts": watchlist.get("counts", {}),
            "summary": watchlist.get("summary", []),
        },
        "model_workflow_health": {
            "model_monitoring": {
                "summary": model_monitoring.get("summary", []),
                "latest_drift_observation": model_monitoring.get("latest_drift_observation"),
                "latest_champion_challenger": model_monitoring.get("latest_champion_challenger"),
            },
            "model_registry": {
                "production_version": model_ops.get("production_version"),
                "deployed_version": (model_ops.get("runtime") or {}).get("model_version"),
                "approval_pending_count": model_ops.get("approval_pending_count"),
                "latest_deployments": (model_ops.get("deployment_history") or [])[:5],
            },
            "business_impact": {
                "summary": scorer_outcomes.get("summary", []),
                "impact_summary": scorer_outcomes.get("impact_summary", []),
                "version_count": scorer_outcomes.get("totals", {}).get("version_count"),
                "score_count": scorer_outcomes.get("totals", {}).get("score_count"),
            },
            "workflow": {
                "counts": workflow.get("counts", {}),
                "summary": workflow.get("summary", []),
                "n8n": (workflow.get("n8n") or {}).get("counts", {}),
                "camunda": (workflow.get("camunda") or {}).get("counts", {}),
            },
        },
        "team_playbook_trends": team_trends,
        "region_playbook_trends": region_trends,
        "step_heatmap_trends": heatmap_trends,
        "playbook_effectiveness": playbook_effectiveness[:16],
        "false_positive_by_team": _breakdown_rows(false_positive_team),
        "false_positive_by_typology": _breakdown_rows(false_positive_typology),
        "false_positive_by_owner": _breakdown_rows(false_positive_owner),
        "case_to_sar_by_typology": case_to_sar_by_typology[:12],
        "filed_sar_volume_by_team": filed_by_team_rows,
        "filed_sar_volume_by_region": filed_by_region_rows,
        "backlog_aging_trends": backlog_trends,
        "compliance_posture": compliance_posture,
        "outcome_correlations": outcome_correlation_rows[:18],
        "workflow_effectiveness": workflow_effectiveness,
        "decision_quality": decision_quality,
        "period_over_period": period_over_period,
        "summary": summary,
    }
    reporting_settings = await get_reporting_automation_settings()
    reporting_alerts, action_recommendations = _evaluate_reporting_alerts(
        report=report_stub,
        period_over_period=period_over_period,
        settings_payload=reporting_settings,
    )
    action_recommendations = list(action_recommendations) + list((decision_quality or {}).get("quality_tuning_recommendations") or [])
    report_stub["reporting_alerts"] = reporting_alerts
    report_stub["action_recommendations"] = action_recommendations
    report_stub["board_reporting"] = _build_board_reporting(report=report_stub, period_over_period=period_over_period)

    return report_stub


async def capture_reporting_snapshot(
    *,
    actor: str | None,
    snapshot_scope: str = "manager",
    snapshot_granularity: str = "daily",
    range_days: int = 180,
    source: str = "manual",
    reference_time: datetime | None = None,
    metadata: dict[str, Any] | None = None,
) -> dict[str, Any]:
    period_start, period_end, period_label = _period_bounds(snapshot_granularity, reference_time)
    report = await get_management_reporting_overview(range_days=range_days, snapshot_scope=snapshot_scope)
    summary_metrics = _build_snapshot_metrics(report)
    summary_metrics.update({
        "period_label": period_label,
        "snapshot_granularity": snapshot_granularity,
    })
    pool = get_pool()
    async with pool.acquire() as conn:
        row = await conn.fetchrow(
            """
            INSERT INTO reporting_snapshots (
                snapshot_scope, snapshot_granularity, range_days, period_start, period_end, period_label, captured_by, source, summary_metrics, snapshot, metadata
            )
            VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9::jsonb,$10::jsonb,$11::jsonb)
            RETURNING id, captured_at, period_start, period_end, period_label
            """,
            snapshot_scope,
            snapshot_granularity,
            range_days,
            period_start,
            period_end,
            period_label,
            actor,
            source,
            json.dumps(summary_metrics),
            json.dumps(jsonable_encoder(report)),
            json.dumps((metadata or {}) | {"period_label": period_label}),
        )
    return {
        "captured": True,
        "snapshot_scope": snapshot_scope,
        "snapshot_granularity": snapshot_granularity,
        "range_days": range_days,
        "snapshot_id": row["id"],
        "period_start": row["period_start"],
        "period_end": row["period_end"],
        "period_label": row["period_label"],
        "captured_at": row["captured_at"],
        "summary": [
            f"Captured {snapshot_scope} {snapshot_granularity} reporting snapshot for {range_days} days.",
            f"Active backlog {summary_metrics.get('active_backlog', 0)} · filed SARs {summary_metrics.get('filed_sars', 0)}.",
        ],
    }


async def get_reporting_snapshots(
    *,
    snapshot_scope: str = "manager",
    snapshot_granularity: str = "daily",
    range_days: int = 180,
    limit: int = 90,
    auto_capture: bool = False,
) -> dict[str, Any]:
    pool = get_pool()
    async with pool.acquire() as conn:
        rows = await conn.fetch(
            """
            SELECT id, snapshot_scope, snapshot_granularity, range_days, period_start, period_end, period_label,
                   captured_at, captured_by, source, summary_metrics, metadata
            FROM reporting_snapshots
            WHERE snapshot_scope = $1 AND snapshot_granularity = $2 AND range_days = $3
            ORDER BY period_end DESC NULLS LAST, captured_at DESC
            LIMIT $4
            """,
            snapshot_scope,
            snapshot_granularity,
            range_days,
            limit,
        )
    if not rows and auto_capture:
        await capture_reporting_snapshot(
            actor="reporting-auto-capture",
            snapshot_scope=snapshot_scope,
            snapshot_granularity=snapshot_granularity,
            range_days=range_days,
            source="auto_bootstrap",
            metadata={"auto_capture": True},
        )
        async with pool.acquire() as conn:
            rows = await conn.fetch(
                """
                SELECT id, snapshot_scope, snapshot_granularity, range_days, period_start, period_end, period_label,
                       captured_at, captured_by, source, summary_metrics, metadata
                FROM reporting_snapshots
                WHERE snapshot_scope = $1 AND snapshot_granularity = $2 AND range_days = $3
                ORDER BY period_end DESC NULLS LAST, captured_at DESC
                LIMIT $4
                """,
                snapshot_scope,
                snapshot_granularity,
                range_days,
                limit,
            )
    points = []
    for row in rows:
        points.append(
            {
                "id": row["id"],
                "snapshot_scope": row["snapshot_scope"],
                "snapshot_granularity": row["snapshot_granularity"] or "daily",
                "range_days": row["range_days"],
                "period_start": row["period_start"],
                "period_end": row["period_end"],
                "period_label": row["period_label"],
                "captured_at": row["captured_at"],
                "captured_by": row["captured_by"],
                "source": row["source"],
                "summary_metrics": _normalize_json_dict(row["summary_metrics"]),
                "metadata": _normalize_json_dict(row["metadata"]),
            }
        )
    latest = points[0] if points else None
    period_over_period = _compute_period_over_period(points)
    return {
        "generated_at": _utcnow(),
        "snapshot_scope": snapshot_scope,
        "snapshot_granularity": snapshot_granularity,
        "range_days": range_days,
        "total_points": len(points),
        "points": points,
        "period_over_period": period_over_period,
        "summary": [
            f"{len(points)} persisted {snapshot_granularity} reporting snapshot(s) available for {snapshot_scope}.",
            (
                f"Latest period {latest.get('period_label') or latest['captured_at'].isoformat()} captured with backlog {latest['summary_metrics'].get('active_backlog', 0)}."
                if latest
                else "No reporting snapshots captured yet."
            ),
        ],
    }


async def get_reporting_drilldown(
    *,
    metric_key: str,
    range_days: int = 180,
    snapshot_id: str | None = None,
    typology: str | None = None,
    team_key: str | None = None,
    region_key: str | None = None,
    owner: str | None = None,
    priority: str | None = None,
    feedback_key: str | None = None,
    limit: int = 40,
) -> dict[str, Any]:
    metric = str(metric_key or "active_backlog").strip().lower()
    if metric.startswith("decision_"):
        return await get_decision_quality_drilldown(
            metric_key=metric,
            range_days=range_days,
            typology=typology,
            team_key=team_key,
            region_key=region_key,
            feedback_key=feedback_key,
            limit=limit,
        )

    since = None
    until = None
    selected_snapshot = None
    if snapshot_id:
        selected_snapshot = await _load_snapshot_by_id(snapshot_id)
        if selected_snapshot:
            since = selected_snapshot.get("period_start")
            until = selected_snapshot.get("period_end")
            range_days = int(selected_snapshot.get("range_days") or range_days)
    case_rows, _alert_rows = await _fetch_reporting_rows(range_days, since=since, until=until)
    normalized_typology = str(typology or "").strip().lower() or None
    normalized_team = str(team_key or "").strip().lower() or None
    normalized_region = str(region_key or "").strip().lower() or None
    normalized_owner = str(owner or "").strip().lower() or None
    normalized_priority = str(priority or "").strip().lower() or None
    metric = str(metric_key or "active_backlog").strip().lower()

    cases: list[dict[str, Any]] = []
    drill_path = ["KPI", metric.replace("_", " ").title()]
    if normalized_typology:
        drill_path.append(_titleize(normalized_typology))
    if normalized_team:
        drill_path.append(_titleize(normalized_team))
    if normalized_region:
        drill_path.append(_titleize(normalized_region))
    if selected_snapshot:
        drill_path.append(str(selected_snapshot.get("period_label") or "Selected Period"))
    for case in case_rows:
        scope = _case_scope(case)
        metadata = _normalize_json_dict(case.get("metadata"))
        playbook_meta = _normalize_json_dict(metadata.get("playbook"))
        case_typology = str(playbook_meta.get("typology") or metadata.get("typology") or ((case.get("alert_types") or [None])[0] or "unknown")).lower()
        case_priority = str(case.get("priority") or "medium").lower()
        assigned_to = str(case.get("assigned_to") or "").lower() or None
        if normalized_typology and case_typology != normalized_typology:
            continue
        if normalized_team and str(scope["team_key"]).lower() != normalized_team:
            continue
        if normalized_region and str(scope["region_key"]).lower() != normalized_region:
            continue
        if normalized_owner and assigned_to != normalized_owner:
            continue
        if normalized_priority and case_priority != normalized_priority:
            continue

        alert_statuses = [str(item).lower() for item in (case.get("alert_statuses") or [])]
        has_sar = bool(case.get("sar_id"))
        is_filed = str(case.get("status") or "").lower() == "sar_filed"
        is_false_positive = "false_positive" in alert_statuses
        created_at = _safe_datetime(case.get("created_at"))
        updated_at = _safe_datetime(case.get("updated_at"))
        workflow_delay_hours = _hours_between(created_at, updated_at)
        drafted_at = _safe_datetime(case.get("drafted_at"))
        reviewed_at = _safe_datetime(case.get("reviewed_at"))
        approved_at = _safe_datetime(case.get("approved_at"))
        filed_at = _safe_datetime(case.get("filed_at"))
        review_lag_hours = _hours_between(drafted_at or created_at, reviewed_at)
        approval_lag_hours = _hours_between(reviewed_at or drafted_at or created_at, approved_at)
        filing_lag_hours = _hours_between(approved_at or reviewed_at or drafted_at or created_at, filed_at)
        evidence_count = int(case.get("evidence_count") or 0)
        included_evidence_count = int(case.get("included_evidence_count") or 0)
        required_evidence_min = 2 if has_sar or str(case.get("status") or "").lower() in {"pending_sar", "sar_filed"} else 1
        missing_evidence_count = max(0, required_evidence_min - included_evidence_count)
        expected_events = 1 + (1 if has_sar else 0) + (1 if str(case.get("sar_status") or "") in {"pending_review", "approved", "filed"} else 0) + (1 if str(case.get("sar_status") or "") in {"approved", "filed"} else 0) + (1 if is_filed else 0)
        present_events = 1 + (1 if case.get("has_sar_drafted_event") else 0) + (1 if case.get("has_sar_submitted_event") else 0) + (1 if case.get("has_sar_approved_event") else 0) + (1 if case.get("has_sar_filed_event") else 0)
        audit_trail_completeness = round(min(1.0, present_events / max(1, expected_events)), 4)
        evidence_pack_completeness = round(min(1.0, included_evidence_count / max(1, required_evidence_min)), 4)

        include = False
        if metric in {"active_backlog", "backlog_aging", "breached_sars"}:
            include = str(case.get("status") or "").lower() not in {"sar_filed", "closed"}
        elif metric in {"typology_mix", "case_to_sar_conversion"}:
            include = True
        elif metric == "filed_sar_volume":
            include = is_filed
        elif metric == "false_positive_rate":
            include = is_false_positive
        elif metric == "review_lag":
            include = review_lag_hours is not None
        elif metric == "approval_lag":
            include = approval_lag_hours is not None
        elif metric == "filing_timeliness":
            include = filing_lag_hours is not None
        elif metric == "audit_trail_completeness":
            include = audit_trail_completeness < 1.0
        elif metric == "evidence_pack_completeness":
            include = evidence_pack_completeness < 1.0
        else:
            include = True

        if not include:
            continue

        cases.append(
            {
                "case_id": case["id"],
                "case_ref": case.get("case_ref"),
                "status": case.get("status"),
                "priority": case.get("priority"),
                "assigned_to": case.get("assigned_to"),
                "team_key": scope["team_key"],
                "team_label": scope["team_label"],
                "region_key": scope["region_key"],
                "region_label": scope["region_label"],
                "typology": case_typology,
                "model_version": _extract_model_version(case.get("alert_rule_ids") or []),
                "sar_status": case.get("sar_status"),
                "has_sar": has_sar,
                "false_positive": is_false_positive,
                "workflow_delay_hours": workflow_delay_hours,
                "review_lag_hours": review_lag_hours,
                "approval_lag_hours": approval_lag_hours,
                "filing_lag_hours": filing_lag_hours,
                "audit_trail_completeness": audit_trail_completeness,
                "evidence_pack_completeness": evidence_pack_completeness,
                "evidence_count": evidence_count,
                "included_evidence_count": included_evidence_count,
                "missing_evidence_count": missing_evidence_count,
                "progress": float(playbook_meta.get("checklist_progress") or 0.0),
                "created_at": created_at,
                "case_deep_link": f"/#case-command?case={case['id']}",
                "recommended_desk": "sar-queue" if has_sar else "manager-console",
            }
        )

    cases.sort(key=lambda item: ((item.get("workflow_delay_hours") or 0), item.get("case_ref") or ""), reverse=True)
    visible = cases[:limit]
    return {
        "generated_at": _utcnow(),
        "metric_key": metric,
        "filters": {
            "range_days": range_days,
            "typology": normalized_typology,
            "team_key": normalized_team,
            "region_key": normalized_region,
            "owner": normalized_owner,
            "priority": normalized_priority,
            "limit": limit,
        },
        "counts": {
            "matched_cases": len(cases),
            "returned_cases": len(visible),
            "sar_cases": sum(1 for item in cases if item.get("has_sar")),
            "false_positive_cases": sum(1 for item in cases if item.get("false_positive")),
        },
        "drill_path": drill_path,
        "snapshot_id": selected_snapshot.get("id") if selected_snapshot else None,
        "period_label": selected_snapshot.get("period_label") if selected_snapshot else None,
        "summary": [
            f"{len(cases)} case(s) matched drilldown metric {metric}.",
            f"Returned {len(visible)} case(s) after applying the UI limit.",
        ],
        "cases": visible,
    }


async def list_report_distribution_rules() -> list[dict[str, Any]]:
    await _ensure_distribution_rules()
    pool = get_pool()
    async with pool.acquire() as conn:
        rows = await conn.fetch("SELECT * FROM report_distribution_rules ORDER BY cadence ASC, display_name ASC")
    return [
        {
            "id": row["id"],
            "rule_key": row["rule_key"],
            "display_name": row["display_name"],
            "target_roles": _normalize_json_list(row["target_roles"]),
            "template_key": row["template_key"],
            "export_format": row["export_format"],
            "cadence": row["cadence"],
            "enabled": bool(row["enabled"]),
            "channels": _normalize_json_list(row["channels"]),
            "recipients": _normalize_json_list(row["recipients"]),
            "metadata": _normalize_json_dict(row["metadata"]),
            "updated_by": row["updated_by"],
            "created_at": row["created_at"],
            "updated_at": row["updated_at"],
        }
        for row in rows
    ]


async def update_report_distribution_rule(*, rule_key: str, actor: str, updates: dict[str, Any]) -> dict[str, Any]:
    await _ensure_distribution_rules()
    pool = get_pool()
    async with pool.acquire() as conn:
        row = await conn.fetchrow("SELECT * FROM report_distribution_rules WHERE rule_key = $1", rule_key)
        if not row:
            raise ValueError(f"Unknown report distribution rule: {rule_key}")
        current = {
            "display_name": row["display_name"],
            "target_roles": _normalize_json_list(row["target_roles"]),
            "template_key": row["template_key"],
            "export_format": row["export_format"],
            "cadence": row["cadence"],
            "enabled": bool(row["enabled"]),
            "channels": _normalize_json_list(row["channels"]),
            "recipients": _normalize_json_list(row["recipients"]),
            "metadata": _normalize_json_dict(row["metadata"]),
        }
        merged = current | {k: v for k, v in updates.items() if v is not None}
        updated = await conn.fetchrow(
            """
            UPDATE report_distribution_rules
            SET display_name = $2,
                target_roles = $3::jsonb,
                template_key = $4,
                export_format = $5,
                cadence = $6,
                enabled = $7,
                channels = $8::jsonb,
                recipients = $9::jsonb,
                metadata = $10::jsonb,
                updated_by = $11,
                updated_at = NOW()
            WHERE rule_key = $1
            RETURNING *
            """,
            rule_key,
            merged["display_name"],
            json.dumps(merged["target_roles"]),
            merged["template_key"],
            merged["export_format"],
            merged["cadence"],
            bool(merged["enabled"]),
            json.dumps(merged["channels"]),
            json.dumps(merged["recipients"]),
            json.dumps(merged["metadata"]),
            actor,
        )
    return {
        "id": updated["id"],
        "rule_key": updated["rule_key"],
        "display_name": updated["display_name"],
        "target_roles": _normalize_json_list(updated["target_roles"]),
        "template_key": updated["template_key"],
        "export_format": updated["export_format"],
        "cadence": updated["cadence"],
        "enabled": bool(updated["enabled"]),
        "channels": _normalize_json_list(updated["channels"]),
        "recipients": _normalize_json_list(updated["recipients"]),
        "metadata": _normalize_json_dict(updated["metadata"]),
        "updated_by": updated["updated_by"],
        "created_at": updated["created_at"],
        "updated_at": updated["updated_at"],
    }


def _report_to_json_bytes(report: dict[str, Any], template_key: str = "manager") -> bytes:
    return json.dumps(jsonable_encoder(report), indent=2, default=_json_default).encode("utf-8")


def _report_sections_for_template(report: dict[str, Any], template_key: str) -> list[tuple[str, list[dict[str, Any]]]]:
    base_sections = [
        ("high_level_kpis", report.get("high_level_kpis", [])),
        ("typology_mix", report.get("typology_mix", [])),
        ("false_positive_by_typology", report.get("false_positive_by_typology", [])),
        ("case_to_sar_by_typology", report.get("case_to_sar_by_typology", [])),
        ("filed_sar_volume_by_team", report.get("filed_sar_volume_by_team", [])),
        ("filed_sar_volume_by_region", report.get("filed_sar_volume_by_region", [])),
        ("backlog_aging_trends", report.get("backlog_aging_trends", [])),
        ("compliance_kpis", (report.get("compliance_posture") or {}).get("kpis", [])),
        ("compliance_by_team", (report.get("compliance_posture") or {}).get("by_team", [])),
        ("compliance_by_region", (report.get("compliance_posture") or {}).get("by_region", [])),
        ("outcome_correlations", report.get("outcome_correlations", [])),
    ]
    if template_key == "executive":
        return [section for section in base_sections if section[0] in {"high_level_kpis", "typology_mix", "filed_sar_volume_by_team", "filed_sar_volume_by_region", "outcome_correlations"}]
    if template_key == "compliance":
        return [section for section in base_sections if section[0] in {"high_level_kpis", "case_to_sar_by_typology", "filed_sar_volume_by_team", "filed_sar_volume_by_region", "backlog_aging_trends", "compliance_kpis", "compliance_by_team", "compliance_by_region", "outcome_correlations"}]
    if template_key == "board":
        return [
            ("high_level_kpis", report.get("high_level_kpis", [])),
            ("typology_mix", report.get("typology_mix", [])),
            ("filed_sar_volume_by_team", report.get("filed_sar_volume_by_team", [])),
            ("outcome_correlations", report.get("outcome_correlations", [])),
            ("board_top_risks", (report.get("board_reporting") or {}).get("top_risks", [])),
            ("board_top_typologies", (report.get("board_reporting") or {}).get("top_typologies", [])),
            ("board_improvements", (report.get("board_reporting") or {}).get("biggest_improvements", [])),
            ("board_declines", (report.get("board_reporting") or {}).get("biggest_declines", [])),
        ]
    return base_sections


def _report_to_csv_bytes(report: dict[str, Any], template_key: str = "manager") -> bytes:
    buffer = StringIO()
    writer = csv.writer(buffer)
    writer.writerow(["section", "key", "label", "value", "note"])
    for section, rows in _report_sections_for_template(report, template_key):
        for item in rows:
            key = item.get("key") or item.get("scope_key") or item.get("typology") or item.get("bucket") or item.get("alert_key") or item.get("recommendation_key") or ""
            label = item.get("label") or item.get("scope_label") or item.get("display_name") or item.get("step_label") or item.get("title") or ""
            value = item.get("value")
            if value is None:
                value = item.get("metric")
            if value is None:
                value = item.get("count")
            if value is None:
                value = item.get("current_value")
            note = item.get("note")
            if note is None and item.get("secondary_metric") is not None:
                note = item.get("secondary_metric")
            if note is None and item.get("secondary_value") is not None:
                note = item.get("secondary_value")
            if note is None:
                note = item.get("message") or item.get("rationale")
            writer.writerow([section, key, label, value, note or ""])
    for line in report.get("monthly_summary", []):
        writer.writerow(["monthly_summary", "", "", line, ""])
    return buffer.getvalue().encode("utf-8")


def _report_to_pdf_bytes(report: dict[str, Any], template_key: str = "manager") -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=0.6 * inch, rightMargin=0.6 * inch, topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    body = ParagraphStyle("ReportBody", parent=styles["BodyText"], fontSize=9.5, leading=13)
    story: list[Any] = []

    def heading(text: str) -> None:
        story.append(Paragraph(text, styles["Heading2"]))
        story.append(Spacer(1, 0.1 * inch))

    def para(text: str) -> None:
        story.append(Paragraph(str(text).replace("\n", "<br/>"), body))
        story.append(Spacer(1, 0.08 * inch))

    story.append(Paragraph(_render_template_title(template_key), styles["Title"]))
    story.append(Spacer(1, 0.12 * inch))
    para(f"Generated at {report.get('generated_at')}")
    para(f"Reporting window: {report.get('range_days')} days")

    heading("Executive Summary")
    for line in report.get("monthly_summary", []):
        para(f"• {line}")

    heading("High-Level KPIs")
    kpi_rows = [["KPI", "Value", "Note"]]
    for item in report.get("high_level_kpis", []):
        kpi_rows.append([str(item.get("label") or ""), str(item.get("value") or ""), str(item.get("note") or "")])
    kpi_table = Table(kpi_rows, repeatRows=1, colWidths=[2.2 * inch, 1.1 * inch, 3.0 * inch])
    kpi_table.setStyle(
        TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eefb")),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#d0d8e8")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7f9fc")]),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
        ])
    )
    story.append(kpi_table)
    story.append(Spacer(1, 0.15 * inch))

    section_map = {
        "typology_mix": "Typology Mix",
        "false_positive_by_typology": "False Positive by Typology",
        "filed_sar_volume_by_team": "Filed SAR Volume by Team",
        "filed_sar_volume_by_region": "Filed SAR Volume by Region",
        "case_to_sar_by_typology": "Case to SAR Conversion",
        "compliance_kpis": "Compliance Oversight KPIs",
        "compliance_by_team": "Compliance by Team",
        "compliance_by_region": "Compliance by Region",
        "outcome_correlations": "Outcome Correlations",
        "board_top_risks": "Top Risks",
        "board_top_typologies": "Top Typologies",
        "board_improvements": "Biggest Improvements",
        "board_declines": "Biggest Declines",
    }
    for section_key, rows in _report_sections_for_template(report, template_key):
        if section_key not in section_map:
            continue
        section_title = section_map[section_key]
        heading(section_title)
        if not rows:
            para("No data available.")
            continue
        for item in rows[:8]:
            label = item.get("scope_label") or item.get("display_name") or item.get("label") or item.get("typology") or "Item"
            value = item.get("metric")
            if value is None:
                value = item.get("count")
            if value is None and item.get("filed_sar_rate") is not None:
                value = item.get("filed_sar_rate")
            para(f"{label}: {value}")

    heading("Health Summary")
    workflow = report.get("model_workflow_health", {}).get("workflow", {})
    monitoring = report.get("model_workflow_health", {}).get("model_monitoring", {})
    para(f"Workflow summary: {' '.join(_as_text_list(workflow.get('summary'))[:2])}")
    para(f"Model monitoring summary: {' '.join(_as_text_list(monitoring.get('summary'))[:2])}")

    doc.build(story)
    return buffer.getvalue()


def _report_to_docx_bytes(report: dict[str, Any], template_key: str = "manager") -> bytes:
    document = Document()
    document.add_heading(_render_template_title(template_key), level=0)
    document.add_paragraph(f"Generated at {report.get('generated_at')}")
    document.add_paragraph(f"Reporting window: {report.get('range_days')} days")
    document.add_heading("Executive Summary", level=1)
    for line in report.get("monthly_summary", []):
        document.add_paragraph(str(line), style="List Bullet")
    document.add_heading("High-Level KPIs", level=1)
    for item in report.get("high_level_kpis", []):
        document.add_paragraph(f"{item.get('label')}: {item.get('value')} ({item.get('note') or 'no note'})")
    document.add_heading("Typology Mix", level=1)
    for item in report.get("typology_mix", []):
        document.add_paragraph(f"{item.get('scope_label')}: {item.get('count')}", style="List Bullet")
    if template_key != "executive":
        document.add_heading("False Positive by Typology", level=1)
        for item in report.get("false_positive_by_typology", []):
            document.add_paragraph(f"{item.get('scope_label')}: {round(float(item.get('metric') or 0) * 100)}%", style="List Bullet")
    if template_key == "compliance":
        document.add_heading("Compliance Oversight KPIs", level=1)
        for item in (report.get("compliance_posture") or {}).get("kpis", []):
            document.add_paragraph(f"{item.get('label')}: {item.get('value')}", style="List Bullet")
        document.add_heading("Compliance Summary", level=1)
        for line in (report.get("compliance_posture") or {}).get("summary", []):
            document.add_paragraph(str(line), style="List Bullet")
    document.add_heading("Filed SAR Volume by Team", level=1)
    for item in report.get("filed_sar_volume_by_team", []):
        document.add_paragraph(f"{item.get('scope_label')}: {item.get('count')}", style="List Bullet")
    if template_key == "compliance":
        document.add_heading("Outcome Correlations", level=1)
        for item in report.get("outcome_correlations", [])[:10]:
            document.add_paragraph(
                f"{item.get('display_name')} / {item.get('priority')}: model v{item.get('dominant_model_version') or 'n/a'}, "
                f"SAR {round(float(item.get('sar_conversion_rate') or 0) * 100)}%, filed {round(float(item.get('filed_sar_rate') or 0) * 100)}%",
                style="List Bullet",
            )
    if template_key == "board":
        document.add_heading("Board Commentary", level=1)
        for line in (report.get("board_reporting") or {}).get("commentary", []):
            document.add_paragraph(str(line), style="List Bullet")
        document.add_heading("Top Risks", level=1)
        for item in (report.get("board_reporting") or {}).get("top_risks", [])[:6]:
            document.add_paragraph(f"{item.get('title')}: {item.get('message') or item.get('metric_label')}", style="List Bullet")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def build_management_report_export(*, range_days: int = 180, snapshot_id: str | None = None, export_format: str = "json", template_key: str = "manager") -> tuple[str, bytes, str]:
    period_suffix = None
    if snapshot_id:
        snapshot = await _load_snapshot_by_id(snapshot_id)
        if not snapshot:
            raise ValueError(f"Unknown reporting snapshot: {snapshot_id}")
        report = snapshot.get("snapshot") or {}
        period_suffix = str(snapshot.get("period_label") or "").replace(" ", "_").replace(":", "-")
    else:
        report = await get_management_reporting_overview(range_days=range_days, snapshot_scope=template_key)
    format_key = str(export_format or "json").lower()
    report_key = f"goaml-{template_key}-report" + (f"-{period_suffix}" if period_suffix else "")
    if format_key == "json":
        return management_report_filename(report_key, "json"), _report_to_json_bytes(report, template_key), "application/json"
    if format_key == "csv":
        return management_report_filename(report_key, "csv"), _report_to_csv_bytes(report, template_key), "text/csv"
    if format_key == "pdf":
        return management_report_filename(report_key, "pdf"), _report_to_pdf_bytes(report, template_key), "application/pdf"
    if format_key == "docx":
        return management_report_filename(report_key, "docx"), _report_to_docx_bytes(report, template_key), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    raise ValueError(f"Unsupported export format: {export_format}")


async def _record_report_notification(
    *,
    notification_type: str,
    channel: str,
    severity: str,
    status: str,
    subject: str,
    target: str | None,
    metadata: dict[str, Any],
) -> None:
    pool = get_pool()
    delivered_at = datetime.now(timezone.utc) if status == "sent" else None
    async with pool.acquire() as conn:
        await conn.execute(
            """
            INSERT INTO notification_events (
                notification_type, channel, severity, status, subject, target, metadata, delivered_at
            )
            VALUES ($1,$2,$3,$4,$5,$6,$7::jsonb,$8)
            """,
            notification_type,
            channel,
            severity,
            status,
            subject,
            target,
            json.dumps(metadata),
            delivered_at,
        )


async def _deliver_report_email(*, subject: str, message: str, recipients: list[str], attachment_name: str, attachment_bytes: bytes, media_type: str) -> tuple[str, str | None]:
    clean_recipients = [item.strip() for item in recipients if str(item).strip()]
    if not (settings.SMTP_HOST and settings.SMTP_FROM and clean_recipients):
        return "not_configured", ", ".join(clean_recipients) if clean_recipients else None
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = settings.SMTP_FROM
    msg["To"] = ", ".join(clean_recipients)
    msg.set_content(message)
    maintype, subtype = media_type.split("/", 1)
    msg.add_attachment(attachment_bytes, maintype=maintype, subtype=subtype, filename=attachment_name)
    with smtplib.SMTP(settings.SMTP_HOST, settings.SMTP_PORT, timeout=20) as smtp:
        if settings.SMTP_USE_TLS:
            smtp.starttls()
        if settings.SMTP_USERNAME and settings.SMTP_PASSWORD:
            smtp.login(settings.SMTP_USERNAME, settings.SMTP_PASSWORD)
        smtp.send_message(msg)
    return "sent", ", ".join(clean_recipients)


async def _deliver_report_slack(*, subject: str, message: str) -> tuple[str, str | None]:
    webhook = settings.SLACK_WEBHOOK_URL
    if not webhook:
        return "not_configured", None
    async with httpx.AsyncClient(timeout=15.0) as client:
        response = await client.post(str(webhook), json={"text": f"{subject}\n{message}"})
        response.raise_for_status()
    return "sent", str(webhook)


async def _deliver_reporting_alert_email(*, subject: str, message: str, recipients: list[str]) -> tuple[str, str | None]:
    clean_recipients = [item.strip() for item in recipients if str(item).strip()]
    if not (settings.SMTP_HOST and settings.SMTP_FROM and clean_recipients):
        return "not_configured", ", ".join(clean_recipients) if clean_recipients else None
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = settings.SMTP_FROM
    msg["To"] = ", ".join(clean_recipients)
    msg.set_content(message)
    with smtplib.SMTP(settings.SMTP_HOST, settings.SMTP_PORT, timeout=20) as smtp:
        if settings.SMTP_USE_TLS:
            smtp.starttls()
        if settings.SMTP_USERNAME and settings.SMTP_PASSWORD:
            smtp.login(settings.SMTP_USERNAME, settings.SMTP_PASSWORD)
        smtp.send_message(msg)
    return "sent", ", ".join(clean_recipients)


async def run_report_distribution(
    *,
    actor: str,
    cadence: str,
    range_days: int = 180,
) -> dict[str, Any]:
    rules = await list_report_distribution_rules()
    selected = [rule for rule in rules if rule.get("enabled") and str(rule.get("cadence") or "").lower() == str(cadence or "daily").lower()]
    results: list[dict[str, Any]] = []
    delivered = 0
    skipped = 0
    for rule in selected:
        effective_range_days = range_days
        if str(rule.get("cadence") or "").lower() == "quarterly":
            effective_range_days = max(range_days, 365)
        filename, payload, media_type = await build_management_report_export(
            range_days=effective_range_days,
            export_format=str(rule.get("export_format") or "pdf"),
            template_key=str(rule.get("template_key") or "manager"),
        )
        subject = f"goAML {str(rule.get('display_name') or 'report')}"
        message = (
            f"Template: {rule.get('template_key')}\n"
            f"Format: {rule.get('export_format')}\n"
            f"Cadence: {rule.get('cadence')}\n"
            f"Generated by: {actor}"
        )
        channel_results: list[dict[str, Any]] = []
        for channel in rule.get("channels") or []:
            if channel == "email":
                status, target = await _deliver_report_email(
                    subject=subject,
                    message=message,
                    recipients=list(rule.get("recipients") or []),
                    attachment_name=filename,
                    attachment_bytes=payload,
                    media_type=media_type,
                )
            elif channel == "slack":
                status, target = await _deliver_report_slack(subject=subject, message=message)
            else:
                status, target = "skipped", None
            if status == "sent":
                delivered += 1
            else:
                skipped += 1
            await _record_report_notification(
                notification_type="board_report_distribution" if str(rule.get("template_key") or "") == "board" else "management_report",
                channel=channel,
                severity="info",
                status=status,
                subject=subject,
                target=target,
                metadata={
                    "rule_key": rule.get("rule_key"),
                    "template_key": rule.get("template_key"),
                    "export_format": rule.get("export_format"),
                    "filename": filename,
                    "triggered_by": actor,
                },
            )
            channel_results.append({"channel": channel, "status": status, "target": target})
        results.append(
            {
                "rule_key": rule.get("rule_key"),
                "display_name": rule.get("display_name"),
                "template_key": rule.get("template_key"),
                "export_format": rule.get("export_format"),
                "channels": rule.get("channels") or [],
                "channel_results": channel_results,
            }
        )
    return {
        "triggered_at": _utcnow(),
        "cadence": cadence,
        "processed_count": len(selected),
        "delivered_count": delivered,
        "skipped_count": skipped,
        "rules": results,
        "summary": [
            f"Processed {len(selected)} distribution rule(s) for cadence {cadence}.",
            f"Delivered {delivered} channel dispatch(es); skipped or not configured {skipped}.",
        ],
    }


async def get_reporting_control_overview(*, range_days: int = 180, snapshot_scope: str = "manager") -> dict[str, Any]:
    overview = await get_management_reporting_overview(range_days=range_days, snapshot_scope=snapshot_scope)
    settings_payload = await get_reporting_automation_settings()
    pool = get_pool()
    async with pool.acquire() as conn:
        alert_rows = await conn.fetch(
            """
            SELECT id, notification_type, channel, severity, status, subject, target, metadata, created_at
            FROM notification_events
            WHERE notification_type IN ('reporting_threshold_alert', 'board_report_distribution')
            ORDER BY created_at DESC
            LIMIT 20
            """
        )
    recent = [{**dict(row), "metadata": _normalize_json_dict(row.get("metadata"))} for row in alert_rows]
    return {
        "generated_at": _utcnow(),
        "snapshot_scope": snapshot_scope,
        "settings": settings_payload,
        "alerts": overview.get("reporting_alerts", []),
        "recommendations": overview.get("action_recommendations", []),
        "board_reporting": overview.get("board_reporting", {}),
        "recent_notifications": recent,
        "summary": [
            f"{len(overview.get('reporting_alerts', []))} active reporting alert(s) for the {snapshot_scope} view.",
            f"{len(overview.get('action_recommendations', []))} manager recommendation(s) are currently available.",
        ],
    }


async def run_reporting_alerts(
    *,
    actor: str,
    channels: list[str],
    snapshot_scope: str = "manager",
    range_days: int = 180,
    force: bool = False,
) -> dict[str, Any]:
    overview = await get_management_reporting_overview(range_days=range_days, snapshot_scope=snapshot_scope)
    alerts = overview.get("reporting_alerts", [])
    recommendations = overview.get("action_recommendations", [])
    if not alerts and not force:
        return {
            "triggered_at": _utcnow(),
            "triggered_by": actor,
            "processed_count": 0,
            "sent_count": 0,
            "skipped_count": 0,
            "alerts": [],
            "recommendations": recommendations,
            "summary": [f"No reporting threshold alerts were triggered for the {snapshot_scope} view."],
        }

    sent = 0
    skipped = 0
    roles = {
        "manager": ["manager"],
        "executive": ["manager", "admin"],
        "compliance": ["auditor", "admin"],
        "board": ["manager", "admin"],
    }.get(snapshot_scope, ["manager"])
    recipients = [rule for rule in await list_report_distribution_rules() if snapshot_scope in {rule.get("template_key"), "manager"}]
    email_targets: list[str] = []
    for rule in recipients:
        email_targets.extend([item for item in rule.get("recipients", []) if item])
    email_targets = sorted(set(email_targets))

    for alert in alerts if alerts else ([{
        "alert_key": "reporting_stability",
        "title": f"{snapshot_scope.title()} reporting stable",
        "severity": "info",
        "status": "stable",
        "metric_key": "reporting_status",
        "metric_label": "Reporting status",
        "current_value": "stable",
        "threshold_value": None,
        "delta_value": None,
        "delta_pct": None,
        "message": f"No reporting thresholds breached for the {snapshot_scope} view.",
        "recommendation_key": None,
        "metadata": {},
    }] if force else []):
        subject = f"goAML reporting alert - {alert.get('title')}"
        deep_link = "/#reports"
        message = (
            f"{alert.get('message')}\n"
            f"Severity: {alert.get('severity')}\n"
            f"Scope: {snapshot_scope}\n"
            f"Metric: {alert.get('metric_label')}\n"
            f"Triggered by: {actor}"
        )
        for channel in channels:
            if channel == "slack":
                status, target = await _deliver_report_slack(subject=subject, message=message)
            elif channel == "email":
                status, target = await _deliver_reporting_alert_email(subject=subject, message=message, recipients=email_targets)
            elif channel == "app":
                status, target = "sent", deep_link
            else:
                status, target = "skipped", None
            if status == "sent":
                sent += 1
            else:
                skipped += 1
            await _record_report_notification(
                notification_type="reporting_threshold_alert",
                channel=channel,
                severity=str(alert.get("severity") or "info"),
                status=status,
                subject=subject,
                target=target,
                metadata={
                    "alert_key": alert.get("alert_key"),
                    "snapshot_scope": snapshot_scope,
                    "recommendation_key": alert.get("recommendation_key"),
                    "roles": roles,
                    "deeplink": deep_link,
                    "message": alert.get("message"),
                    "triggered_by": actor,
                },
            )
    return {
        "triggered_at": _utcnow(),
        "triggered_by": actor,
        "processed_count": len(alerts) if alerts else (1 if force else 0),
        "sent_count": sent,
        "skipped_count": skipped,
        "alerts": alerts,
        "recommendations": recommendations,
        "summary": [
            f"Processed {len(alerts) if alerts else (1 if force else 0)} reporting alert item(s) for the {snapshot_scope} view.",
            f"Sent {sent} channel dispatch(es); skipped or not configured {skipped}.",
        ],
    }
